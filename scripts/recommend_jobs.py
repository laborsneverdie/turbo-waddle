"""
岗位推荐脚本（每8小时由 GitHub Actions 触发）
使用 Supabase REST API，不依赖数据库直连

流程：
  1. 通过 REST API 读取所有用户资料
  2. 从7个招聘数据源爬取真实岗位（国聘网/央企/地方国企/微信/智联/BOSS/国资委）
  3. 通过 REST API 写入 job_recommendations 表
  4. 生成 H5/docx 岗位推荐报告
  5. 通过 PushPlus 微信推送 + SMTP 邮件推送
"""

import json
import time
import re
import os
import sys
import html
import random
import smtplib
import traceback
import urllib.parse
from urllib.parse import quote
from datetime import datetime
from email.mime.text import MIMEText

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ============================================================
# Supabase 客户端初始化
# ============================================================
SUPABASE_URL = os.environ.get("SUPABASE_URL", "https://clihwbzomhctkxooldbz.supabase.co")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY") or os.environ.get("SUPABASE_SERVICE_ROLE_KEY", "")

# ============================================================
# 常量定义
# ============================================================

# Supabase REST API headers
_SUPABASE_HEADERS = {
    "apikey": SUPABASE_KEY,
    "Authorization": f"Bearer {SUPABASE_KEY}",
    "Content-Type": "application/json",
}

# 通用请求头
COMMON_UA = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/125.0.0.0 Safari/537.36"
)

# 请求超时（秒）
REQUEST_TIMEOUT = 20

# Supabase API 数据读写
def fetch_users():
    """从 Supabase REST API 读取所有用户资料"""
    resp = requests.get(
        f"{SUPABASE_URL}/rest/v1/user_profiles?order=created_at.desc",
        headers=_SUPABASE_HEADERS, timeout=30,
    )
    if resp.status_code != 200:
        print(f"[Supabase] 读取用户失败: {resp.status_code} {resp.text}")
        return []
    return resp.json()

def save_recommendations(user_id: int, jobs: list[dict]):
    """通过 REST API 写入 job_recommendations 表"""
    _save_headers = dict(_SUPABASE_HEADERS)
    _save_headers["Prefer"] = "return=minimal"
    
    # 先删除旧推荐
    resp_del = requests.delete(
        f"{SUPABASE_URL}/rest/v1/job_recommendations?user_id=eq.{user_id}",
        headers=_save_headers, timeout=30,
    )
    if resp_del.status_code not in (200, 204):
        print(f"  [Supabase] 删除旧推荐失败: {resp_del.status_code} {resp_del.text[:200]}")
    
    saved_count = 0
    for idx, j in enumerate(jobs):
        row = {
            "user_id": user_id,
            "job_title": j.get("job_title", "")[:500],
            "company": j.get("company", "")[:200],
            "enterprise_type": j.get("enterprise_type", "")[:50],
            "match_score": j.get("match_score", 0),
            "detail_link": j.get("detail_link", "")[:500],
        }
        try:
            resp = requests.post(
                f"{SUPABASE_URL}/rest/v1/job_recommendations",
                headers=_save_headers, json=row, timeout=30,
            )
            if resp.status_code in (200, 201, 204):
                saved_count += 1
            else:
                if idx < 2:  # 只打印前2条的详细错误
                    print(f"  [Supabase] 写入岗位{idx+1}失败: {resp.status_code} {resp.text[:300]}")
        except Exception as e:
            if idx < 2:
                print(f"  [Supabase] 写入岗位{idx+1}异常: {e}")
    if saved_count:
        print(f"  [Supabase] 成功写入 {saved_count}/{len(jobs)} 条岗位")
    else:
        if len(jobs) > 0:
            print(f"  [Supabase] ⚠️ 全部 {len(jobs)} 条写入失败！检查表结构/RLS/字段类型")
    return jobs  # 返回原始完整数据供 H5/docx 使用

# 每个网站爬取超时
SITE_TIMEOUT = 20

# 请求间隔（秒）
REQUEST_INTERVAL_MIN = 1.0
REQUEST_INTERVAL_MAX = 2.0

# 学历等级映射（数值越大学历越高）
EDU_LEVEL = {
    "不限": 0,
    "大专": 1,
    "专科": 1,
    "本科": 2,
    "学士": 2,
    "硕士": 3,
    "研究生": 3,
    "博士": 4,
    "MBA/EMBA": 3,
}

# 经验等级映射（数值越大经验要求越高）
EXP_LEVEL = {
    "不限": 0,
    "无经验": 0,
    "应届": 1,
    "应届生": 1,
    "应届毕业生": 1,
    "1年以下": 1,
    "1年": 1,
    "1-3年": 2,
    "1~3年": 2,
    "3-5年": 3,
    "3~5年": 3,
    "5-10年": 4,
    "5~10年": 4,
    "10年以上": 5,
}

# ============================================================
# 央企招聘官网数据源（URL + 名称）
# ============================================================
NATIONAL_SOE_SOURCES = [
    {
        "name": "中石油",
        "urls": [
            "http://zhaopin.cnpc.com.cn/",
        ],
        "wechat_name": "中国石油招聘",
    },
    {
        "name": "南方电网",
        "urls": [
            "https://zhaopin.csg.cn/",
            "https://zhaopin.csg.cn/#/recruitment-social",
        ],
        "wechat_name": "南网50Hz",
    },
    {
        "name": "国家电网",
        "urls": [
            "https://zhaopin.sgcc.com.cn/",
        ],
        "wechat_name": "国家电网招聘",
    },
    {
        "name": "中石化",
        "urls": [
            "https://job.sinopec.com/",
        ],
        "wechat_name": "中国石化招聘",
    },
    {
        "name": "中国海油",
        "urls": [
            "https://zhaopin.cnooc.com.cn/",
        ],
        "wechat_name": "中国海油",
    },
    {
        "name": "中国移动",
        "urls": [
            "https://job.10086.cn/",
            "https://campus.10086.cn/",
        ],
        "wechat_name": "中国移动招聘",
    },
    {
        "name": "中国电信",
        "urls": [
            "https://campus.51job.com/chinatelecom/",
        ],
        "wechat_name": "中国电信招聘",
    },
    {
        "name": "中国联通",
        "urls": [
            "https://chinaunicom.zhiye.com/",
            "https://hr.chinaunicom.cn/",
        ],
        "wechat_name": "中国联通招聘",
    },
    {
        "name": "中国建筑",
        "urls": [
            "https://cscec.zhiye.com/",
        ],
        "wechat_name": "中国建筑招聘",
    },
    {
        "name": "华润集团",
        "urls": [
            "https://crc.wintalent.cn/",
            "https://www.crc.com.cn/rlzy_43682/zxns/",
        ],
        "wechat_name": "华润招聘",
    },
    {
        "name": "招商局集团",
        "urls": [
            "https://www.cmhk.com/main/rlzy/rczp/",
        ],
        "wechat_name": "百年招商局",
    },
    {
        "name": "中国中铁",
        "urls": [
            "https://www.crecg.com/",
            "https://www.crecg.com.cn/",
        ],
        "wechat_name": "中国中铁招聘",
    },
    {
        "name": "中粮集团",
        "urls": [
            "https://cofco-campus.zhiye.com/",
            "http://campus.51job.com/cofco/",
            "https://cofco.zhiye.com/",
        ],
        "wechat_name": "中粮招聘",
    },
    {
        "name": "中国航天科工",
        "urls": [
            "https://zhaopin.casic.cn/",
        ],
        "wechat_name": "中国航天科工招聘",
    },
    {
        "name": "中国船舶",
        "urls": [
            "https://zhaopin.cssc.net.cn/",
        ],
        "wechat_name": "中国船舶招聘",
    },
    {
        "name": "国家电投",
        "urls": [
            "https://hr.cpic.com.cn/",
        ],
        "wechat_name": "国家电投",
    },
]

# ============================================================
# 城市 → 地方国企映射表
# ============================================================
LOCAL_SOE_MAP = {
    "长沙": [
        {"name": "长沙银行", "urls": ["https://cscb.zhiye.com/", "https://www.cscb.cn/"]},
        {"name": "湖南银行", "urls": ["https://www.hunan-bank.com/96599/gywx/rczp/index.shtml", "https://www.hunan-bank.com/"]},
        {"name": "湖南建投集团", "urls": ["https://www.hncig.cn/channel/29080.html", "https://www.hncig.cn/", "https://www.hnjg.com.cn/"]},
        {"name": "兴湘集团", "urls": ["https://www.hnxtg.com/"]},
        {"name": "湖南高速集团", "urls": ["https://www.hngs.net/"]},
        {"name": "长沙轨交集团", "urls": ["https://www.hncsmtr.com/"]},
        {"name": "湖南中烟工业", "urls": ["http://www.hnti.com.cn/"]},
        {"name": "长沙水业集团", "urls": ["http://www.cssy.com.cn/"]},
        {"name": "湖南机场管理集团", "urls": ["https://www.hna-ca.com.cn/"]},
    ],
    "北京": [
        {"name": "北京银行", "urls": ["https://www.bankofbeijing.com.cn/"]},
        {"name": "北京农商银行", "urls": ["https://www.bjrcb.com/"]},
        {"name": "首钢集团", "urls": ["https://www.shougang.com.cn/"]},
        {"name": "北汽集团", "urls": ["https://www.baicgroup.com.cn/"]},
        {"name": "京能集团", "urls": ["https://www.jnenergy.com/"]},
    ],
    "上海": [
        {"name": "浦发银行", "urls": ["https://www.spdb.com.cn/"]},
        {"name": "上海银行", "urls": ["https://www.bosc.cn/"]},
        {"name": "上汽集团", "urls": ["https://www.saicsa.com/", "https://www.saicmotor.com/"]},
        {"name": "上海电气", "urls": ["https://www.shanghai-electric.com/"]},
        {"name": "申通地铁", "urls": ["https://www.shmetro.com/"]},
    ],
    "深圳": [
        {"name": "深圳地铁", "urls": ["https://www.szmc.net/"]},
        {"name": "深圳能源", "urls": ["https://www.sec.com.cn/"]},
        {"name": "深圳水务", "urls": ["https://www.sz-water.com.cn/"]},
    ],
    "广州": [
        {"name": "广州银行", "urls": ["https://www.gzcb.com.cn/"]},
        {"name": "广汽集团", "urls": ["https://www.gac.com.cn/"]},
        {"name": "广州地铁", "urls": ["https://www.gzmtr.com/"]},
    ],
    "成都": [
        {"name": "成都银行", "urls": ["https://www.bocd.com.cn/"]},
        {"name": "成都轨交集团", "urls": ["https://www.chengdurail.com/"]},
    ],
    "武汉": [
        {"name": "武汉农村商业银行", "urls": ["https://www.whrcb.com/"]},
        {"name": "武汉地铁", "urls": ["https://www.wuhanrt.com/"]},
        {"name": "湖北交投集团", "urls": ["https://www.hbjt.com.cn/"]},
    ],
    "杭州": [
        {"name": "杭州银行", "urls": ["https://www.hzbank.com.cn/"]},
        {"name": "杭州地铁", "urls": ["https://www.hzmetro.com/"]},
    ],
    "南京": [
        {"name": "南京银行", "urls": ["https://www.njcb.com.cn/"]},
        {"name": "南京地铁", "urls": ["https://www.njmetro.com.cn/"]},
    ],
    "西安": [
        {"name": "长安银行", "urls": ["https://www.ccbbchina.com/"]},
        {"name": "西安地铁", "urls": ["http://www.xianrail.com/"]},
    ],
    "重庆": [
        {"name": "重庆银行", "urls": ["https://www.cqcbank.com/"]},
        {"name": "重庆农村商业银行", "urls": ["https://www.cqrcb.com/"]},
        {"name": "重庆轨交集团", "urls": ["https://www.cqmetro.cn/"]},
    ],
}


# 同省城市映射（用于城市匹配的扩展判断）
PROVINCE_MAP = {
    "长沙": "湖南", "株洲": "湖南", "湘潭": "湖南", "衡阳": "湖南",
    "武汉": "湖北", "宜昌": "湖北", "襄阳": "湖北",
    "广州": "广东", "深圳": "广东", "佛山": "广东", "东莞": "广东",
    "杭州": "浙江", "宁波": "浙江", "温州": "浙江",
    "成都": "四川", "绵阳": "四川",
    "南京": "江苏", "苏州": "江苏", "无锡": "江苏",
    "西安": "陕西",
    "重庆": "重庆",
}


# ============================================================
# 辅助函数
# ============================================================

def _safe_request(url, method="GET", headers=None, json_body=None, timeout=REQUEST_TIMEOUT, verify_ssl=True):
    """
    安全的 HTTP 请求封装，带异常处理。
    verify_ssl=False 时跳过 SSL 证书验证（用于国企等旧证书网站）。
    """
    default_headers = {
        "User-Agent": COMMON_UA,
        "Accept": "application/json, text/plain, */*",
        "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
    }
    if headers:
        default_headers.update(headers)

    # 跳过 SSL 验证时抑制警告
    if not verify_ssl:
        import urllib3
        urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

    try:
        if method.upper() == "POST":
            resp = requests.post(
                url, headers=default_headers, json=json_body, timeout=timeout, verify=verify_ssl
            )
        else:
            resp = requests.get(url, headers=default_headers, timeout=timeout, verify=verify_ssl)
        resp.raise_for_status()
        return resp
    except requests.exceptions.Timeout:
        print(f"  [WARNING] 请求超时: {url}")
        return None
    except requests.exceptions.ConnectionError as e:
        # SSL 证书问题时自动重试（跳过验证）
        err_str = str(e)
        if any(kw in err_str for kw in ["SSL", "CERTIFICATE", "UNSAFE_LEGACY"]):
            print(f"  [INFO] SSL证书问题，跳过验证重试: {url}")
            try:
                import urllib3
                urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
                if method.upper() == "POST":
                    resp = requests.post(
                        url, headers=default_headers, json=json_body, timeout=timeout, verify=False
                    )
                else:
                    resp = requests.get(url, headers=default_headers, timeout=timeout, verify=False)
                resp.raise_for_status()
                return resp
            except Exception as e2:
                # ★ SSL验证跳过也失败了 → 尝试启用旧版SSL重协商（OP_LEGACY_SERVER_CONNECT）
                err_s2 = str(e2)
                if any(kw in err_s2 for kw in ["UNSAFE_LEGACY", "SSL", "CERTIFICATE", "sslv3"]):
                    print(f"  [INFO] 旧版SSL服务器，启用LEGACY重协商重试: {url}")
                    try:
                        import ssl as _ssl_module
                        from requests.adapters import HTTPAdapter
                        from urllib3.poolmanager import PoolManager

                        class _LegacySSLAdapter(HTTPAdapter):
                            def init_poolmanager(self, *a, **kw):
                                ctx = _ssl_module.create_default_context()
                                ctx.check_hostname = False
                                ctx.verify_mode = _ssl_module.CERT_NONE
                                # OP_LEGACY_SERVER_CONNECT = 0x4
                                ctx.options |= 0x4
                                kw['ssl_context'] = ctx
                                return super().init_poolmanager(*a, **kw)

                        s = requests.Session()
                        s.mount("https://", _LegacySSLAdapter())
                        s.mount("http://", HTTPAdapter())
                        if method.upper() == "POST":
                            resp = s.post(url, headers=default_headers, json=json_body, timeout=timeout)
                        else:
                            resp = s.get(url, headers=default_headers, timeout=timeout)
                        resp.raise_for_status()
                        s.close()
                        return resp
                    except Exception as e3:
                        print(f"  [WARNING] LegacySSL重试后仍失败: {url} - {e3}")
                        return None
                print(f"  [WARNING] SSL跳过验证后仍失败: {url} - {e2}")
                return None
        print(f"  [WARNING] 连接错误: {url} - {e}")
        return None
    except requests.exceptions.HTTPError as e:
        print(f"  [WARNING] HTTP错误: {url} - {e}")
        return None
    except Exception as e:
        print(f"  [WARNING] 请求异常: {url} - {e}")
        return None


def _parse_edu_level(edu_str):
    """从字符串中解析学历等级"""
    if not edu_str:
        return 0
    edu_str = str(edu_str).strip()
    for key, val in EDU_LEVEL.items():
        if key in edu_str:
            return val
    return 0


def _parse_exp_level(exp_str):
    """从字符串中解析经验等级"""
    if not exp_str:
        return 0
    exp_str = str(exp_str).strip()
    for key, val in EXP_LEVEL.items():
        if key in exp_str:
            return val
    # 尝试数字匹配
    m = re.search(r"(\d+)\s*年", exp_str)
    if m:
        years = int(m.group(1))
        if years <= 1:
            return 1
        elif years <= 3:
            return 2
        elif years <= 5:
            return 3
        elif years <= 10:
            return 4
        else:
            return 5
    return 0


def _parse_experience_years(exp_str):
    """从用户经验字符串中解析年数"""
    if not exp_str:
        return 0
    m = re.search(r"(\d+)", str(exp_str))
    return int(m.group(1)) if m else 0


def _random_delay():
    """随机延迟，避免频繁请求"""
    delay = random.uniform(REQUEST_INTERVAL_MIN, REQUEST_INTERVAL_MAX)
    time.sleep(delay)


def _extract_section(text, keywords, max_lines=10):
    """
    从岗位描述文本中提取特定部分（职责/要求等）。
    策略：找到关键词所在行作为起点，一直捕获到下一个段落标题或文本结尾。
    """
    if not text:
        return ""

    # 所有已知的段落标题关键词
    all_section_headers = [
        "岗位职责", "工作内容", "职位描述", "工作职责", "主要职责", "职位要求",
        "任职要求", "岗位要求", "任职资格", "任职条件", "应聘条件",
        "福利待遇", "薪酬福利", "薪资福利", "工作时间", "工作地点",
        "联系方式", "公司介绍", "企业介绍", "备注",
    ]

    lines = [l.strip() for l in text.split("\n") if l.strip()]
    if not lines:
        return ""

    # 找到关键词起始行
    start_idx = -1
    for i, line in enumerate(lines):
        for kw in keywords:
            if kw in line:
                start_idx = i
                break
        if start_idx >= 0:
            break

    if start_idx < 0:
        return ""

    # 从起始行开始，一直捕获到下一个段落标题（不包括当前段落标题）
    result = []
    for i in range(start_idx, len(lines)):
        line = lines[i]
        # 如果不是起始行，且遇到了新的段落标题，停止
        if i > start_idx:
            is_new_section = any(header in line for header in all_section_headers)
            if is_new_section:
                # 但要排除当前段落的标题在内容中的重复出现
                current_kw_in_line = any(kw in line for kw in keywords)
                if not current_kw_in_line:
                    break
        result.append(line)
        if len(result) >= max_lines:
            break

    return "\n".join(result) if result else ""


def _extract_responsibilities(contents):
    """从岗位描述中提取岗位职责"""
    if not contents:
        return "详见岗位详情页"
    resp = _extract_section(contents, ["岗位职责", "工作内容", "职位描述", "工作职责", "主要职责"])
    if resp:
        return resp
    # 如果没有明确分节，返回前几行
    lines = [l.strip() for l in contents.split("\n") if l.strip()]
    return "\n".join(lines[:8]) if lines else "详见岗位详情页"


def _extract_requirements(contents, edu_str="", exp_str=""):
    """从岗位描述中提取任职要求"""
    if not contents:
        parts = []
        if edu_str:
            parts.append(f"学历要求：{edu_str}")
        if exp_str:
            parts.append(f"经验要求：{exp_str}")
        return "\n".join(parts) if parts else "详见岗位详情页"
    req = _extract_section(contents, ["任职要求", "岗位要求", "任职资格", "任职条件", "应聘条件", "职位要求"])
    if req:
        return req
    # 从全文中提取包含"要求"关键词的行
    lines = contents.split("\n")
    req_lines = [l.strip() for l in lines if any(kw in l for kw in ["要求", "具备", "熟悉", "掌握", "优先", "以上"])]
    if req_lines:
        return "\n".join(req_lines[:8])
    parts = []
    if edu_str:
        parts.append(f"学历要求：{edu_str}")
    if exp_str:
        parts.append(f"经验要求：{exp_str}")
    return "\n".join(parts) if parts else "详见岗位详情页"


def _determine_enterprise_type(company_info_str, source):
    """根据公司信息字符串判断企业类型"""
    text = str(company_info_str) if company_info_str else ""
    if any(kw in text for kw in ["央企", "国企", "国有", "国资委", "中央企业"]):
        return "国企"
    if any(kw in text for kw in ["外资", "外企", "合资", "外商独资"]):
        return "外企"
    if any(kw in text for kw in ["民营", "私企", "私企", "私营", "股份制"]):
        return "私企"
    # 根据来源推断
    if source in ["国聘网", "国资委央企招聘"]:
        return "国企"
    return "私企"


def _generate_development(enterprise_type, source):
    """根据企业类型生成职业发展前景描述"""
    if enterprise_type == "国企":
        return "国企平台稳定，晋升通道清晰，福利保障完善"
    elif enterprise_type == "外企":
        return "国际化平台，职业发展路径多元，薪酬体系完善"
    else:
        return "市场化薪酬，成长空间大，适合快速积累经验"


def _generate_benefits(contents):
    """从岗位描述中提取福利待遇"""
    if not contents:
        return "详见岗位详情页"
    benefit_keywords = [
        "五险一金", "六险二金", "五险二金", "七险一金", "七险二金",
        "年终奖", "年终奖金", "年底双薪",
        "带薪年假", "年假",
        "餐补", "餐费补贴", "免费三餐", "免费午餐",
        "交通补贴", "通讯补贴", "住房补贴",
        "体检", "定期体检",
        "补充医疗", "商业保险",
        "弹性工作", "周末双休", "双休",
        "股票期权", "期权",
        "培训", "晋升", "团建",
    ]
    found = []
    text = str(contents)
    for kw in benefit_keywords:
        if kw in text:
            found.append(kw)
    if found:
        # 去重
        seen = set()
        unique = []
        for b in found:
            if b not in seen:
                seen.add(b)
                unique.append(b)
        return "/".join(unique)
    # 默认福利（国企常见福利）
    return "五险一金/年终奖/带薪年假/定期体检"


def _is_city_match(job_location, user_city):
    """判断岗位地点是否匹配用户城市"""
    if not job_location or not user_city:
        return False
    job_loc = str(job_location).replace("市", "").strip()
    user_c = str(user_city).replace("市", "").strip()
    return user_c in job_loc or job_loc in user_c


# 主要城市列表（用于从文本中提取工作地点）
_ALL_CITIES = [
    "北京", "上海", "广州", "深圳", "杭州", "成都", "长沙", "武汉", "南京",
    "西安", "苏州", "重庆", "天津", "郑州", "青岛", "沈阳", "大连", "厦门",
    "合肥", "济南", "哈尔滨", "福州", "昆明", "贵阳", "南宁", "石家庄",
    "太原", "长春", "南昌", "兰州", "海口", "呼和浩特", "银川", "西宁",
    "乌鲁木齐", "拉萨", "株洲", "湘潭", "衡阳", "宜昌", "襄阳", "佛山",
    "东莞", "宁波", "温州", "绵阳", "无锡", "常州", "绍兴", "嘉兴",
]


def _detect_city_from_text(text):
    """
    从文本中提取城市名称。
    优先匹配 "XX招聘" 或 "【XX招聘】" 模式，其次匹配城市名出现。
    """
    if not text:
        return ""

    text = str(text)

    # 模式1: 【XX招聘】 或 XX招聘
    m = re.search(r"[【\[]?\s*([北京上海广州深圳杭州成都长沙武汉南京西安苏州重庆天津郑州青岛沈阳大连厦门合肥济南哈尔滨福州昆明贵阳南宁石家庄太原长春南昌兰州海口呼和浩特银川西宁乌鲁木齐拉萨株洲湘潭衡阳宜昌襄阳佛山东莞宁波温州绵阳无锡常州绍兴嘉兴]{2,4})\s*(?:招聘|招贤|招新)", text)
    if m:
        return m.group(1)

    # 模式2: 城市名 + 地点关键词
    m = re.search(r"(工作地点|地点|地址|所在地)[：:]\s*([北京上海广州深圳杭州成都长沙武汉南京西安苏州重庆天津郑州青岛沈阳大连厦门合肥济南哈尔滨福州昆明贵阳南宁石家庄太原长春南昌兰州海口呼和浩特银川西宁乌鲁木齐拉萨株洲湘潭衡阳宜昌襄阳佛山东莞宁波温州绵阳无锡常州绍兴嘉兴]{2,4})", text)
    if m:
        return m.group(2)

    # 模式3: 直接匹配已知城市名（取第一个出现的）
    for city in _ALL_CITIES:
        if city in text:
            return city

    return ""


# ============================================================
# 匹配度计算
# ============================================================

def calculate_match_score(job, user):
    """
    计算岗位与用户的匹配度 (0-100)。

    评分维度（城市已在前置筛选完成，专业方向第一优先）：
    - 方向匹配 (0-50)：求职方向与岗位名称/描述相关度（第一优先级）
    - 学历匹配 (0-25)：用户学历 >= 岗位要求 = 高分
    - 经验匹配 (0-25)：用户经验 >= 岗位要求 = 高分
    """
    score = 0

    # --- 方向匹配 (0-50) ★ 第一优先级 ---
    user_field = (user.get("field") or user.get("direction") or "").lower().strip()
    job_title = (job.get("job_title") or "").lower()
    job_desc = (job.get("_raw_contents") or job.get("job_desc") or "").lower()

    if user_field:
        # 拆分：英文=核心技能(Python/Java), 中文=通用角色(开发/工程师)
        en_skills = [w.lower() for w in re.findall(r'[a-zA-Z0-9]+', user_field)]
        cn_roles = re.findall(r'[\u4e00-\u9fff]+', user_field)

        # 生成中文双字词（bigrams），适配"财务会计"→"财务"+"会计"的部分匹配
        cn_bigrams = set()
        for role in cn_roles:
            if len(role) >= 2:
                for i in range(len(role) - 1):
                    cn_bigrams.add(role[i:i+2])

        # 1. 完整字段匹配在标题中 → 满分
        if user_field in job_title:
            score += 50
        # 2. 全部英文技能词出现在标题中
        elif en_skills and all(kw in job_title for kw in en_skills):
            score += 45
        # 3. 至少一个英文技能词出现在标题中 (核心技能对口)
        elif en_skills and any(kw in job_title for kw in en_skills):
            score += 35
        # 4. 英文技能词出现在描述中
        elif en_skills and any(kw in job_desc for kw in en_skills):
            score += 20
        # 5. 中文 bigram 匹配标题 — 纯中文方向的核心逻辑
        elif cn_bigrams:
            title_bigram_match = sum(1 for bg in cn_bigrams if bg in job_title)
            title_bigram_ratio = title_bigram_match / len(cn_bigrams) if cn_bigrams else 0
            if title_bigram_ratio >= 0.5:
                score += 40  # 标题中大部分 bigram 命中
            elif title_bigram_match > 0:
                score += 25  # 标题中部分 bigram 命中
            else:
                # 标题不匹配，看描述
                desc_bigram_match = sum(1 for bg in cn_bigrams if bg in job_desc)
                desc_bigram_ratio = desc_bigram_match / len(cn_bigrams) if cn_bigrams else 0
                if desc_bigram_ratio >= 0.5:
                    score += 15  # 描述中大部分命中
                elif desc_bigram_match > 0:
                    score += 8   # 描述中少量命中
                else:
                    score += 0   # 完全不命中
        # 6. 中文角色词出现在标题中 (短词后备)
        elif cn_roles and any(kw in job_title for kw in cn_roles):
            score += 10
        # 7. 中文角色词出现在描述中
        elif cn_roles and any(kw in job_desc for kw in cn_roles):
            score += 5
        # 7. 完全不对口
        else:
            score += 0
    else:
        score += 25  # 无方向信息，给中间分

    # --- 学历匹配 (0-25) ---
    user_edu = user.get("degree") or user.get("education") or ""
    job_edu = job.get("_raw_edu") or job.get("education_req") or ""
    user_edu_level = _parse_edu_level(user_edu)
    job_edu_level = _parse_edu_level(job_edu)
    if job_edu_level == 0:
        score += 20  # 不限学历
    elif user_edu_level >= job_edu_level:
        score += 25  # 完全满足
    elif user_edu_level == job_edu_level - 1:
        score += 15  # 差一档
    else:
        score += 5

    # --- 经验匹配 (0-25) ---
    user_exp = user.get("experience") or ""
    job_exp = job.get("_raw_exp") or job.get("experience_req") or ""
    user_exp_level = _parse_exp_level(user_exp)
    job_exp_level = _parse_exp_level(job_exp)
    if job_exp_level == 0:
        score += 20  # 不限经验
    elif user_exp_level >= job_exp_level:
        score += 25  # 经验满足
    elif user_exp_level == job_exp_level - 1:
        score += 15  # 差一档
    else:
        score += 5

    # 如果方向完全不匹配，总分上限压低（防止不对口岗位排到前面）
    if user_field and score <= 50:  # 方向0分 + 学历 + 经验 ≤ 50
        score = min(score, 40)

    return min(score, 100)


# ============================================================
# 爬虫：国聘网（API）
# ============================================================

def crawl_iguopin(keyword, city=None):
    """
    爬取国聘网（国资央企招聘平台）岗位数据。
    使用公开 API，最可靠的数据源。
    """
    print(f"\n[国聘网] 开始爬取，关键词: {keyword}，城市: {city}")
    api_url = "https://gp-api.iguopin.com/api/jobs/v1/list"
    headers = {
        "Content-Type": "application/json;charset=UTF-8",
        "Device": "pc",
        "Subsite": "cujiuye",
        "Version": "5.0.0",
        "User-Agent": COMMON_UA,
        "Referer": "https://www.iguopin.com/",
        "Origin": "https://www.iguopin.com",
    }
    payload = {"page": 1, "page_size": 50, "keyword": keyword}

    resp = _safe_request(api_url, method="POST", headers=headers, json_body=payload)
    if not resp:
        print("[国聘网] API请求失败")
        return []

    try:
        data = resp.json()
    except Exception as e:
        print(f"[国聘网] JSON解析失败: {e}")
        return []

    job_list = []
    if isinstance(data, dict):
        job_list = data.get("data", {}).get("list", [])
    elif isinstance(data, list):
        job_list = data

    print(f"[国聘网] API返回 {len(job_list)} 条岗位")

    results = []
    for item in job_list:
        try:
            job_id = item.get("job_id") or item.get("id") or ""
            job_name = item.get("job_name") or item.get("name") or ""
            company_name = item.get("company_name") or item.get("company") or ""

            if not job_name or not company_name:
                continue

            # 薪资
            min_wage = item.get("min_wage") or ""
            max_wage = item.get("max_wage") or ""
            if min_wage and max_wage:
                salary_range = f"{min_wage}-{max_wage}元/月"
            elif min_wage:
                salary_range = f"{min_wage}元/月起"
            else:
                salary_range = item.get("salary") or "面议"

            # 学历/经验
            edu = item.get("education_cn") or item.get("education") or ""
            exp = item.get("experience_cn") or item.get("experience") or item.get("work_years") or ""

            # 地点
            district_list = item.get("district_list") or []
            if isinstance(district_list, list) and district_list:
                work_location = district_list[0].get("area_cn") or ""
            else:
                work_location = item.get("city") or item.get("city_cn") or ""

            # 岗位描述
            contents = item.get("contents") or item.get("description") or item.get("job_desc") or ""

            # 公司信息
            company_info = item.get("company_info") or {}
            if isinstance(company_info, dict):
                nature = company_info.get("nature_cn") or company_info.get("nature") or ""
            else:
                nature = str(company_info)

            enterprise_type = _determine_enterprise_type(nature + " " + company_name, "国聘网")

            detail_link = f"https://www.iguopin.com/job/detail?id={job_id}" if job_id else ""

            job_obj = {
                "job_title": job_name,
                "company": company_name,
                "enterprise_type": enterprise_type,
                "detail_link": detail_link,
                "salary_range": salary_range,
                "responsibilities": _extract_responsibilities(contents),
                "requirements": _extract_requirements(contents, edu, exp),
                "benefits": _generate_benefits(contents),
                "development": _generate_development(enterprise_type, "国聘网"),
                "work_location": work_location or "全国",
                "source": "国聘网（真实数据）",
                "search_keyword": keyword,
                "_raw_contents": contents,
                "_raw_edu": edu,
                "_raw_exp": exp,
            }
            results.append(job_obj)
        except Exception as e:
            print(f"  [WARNING] 解析国聘网岗位失败: {e}")
            continue

    print(f"[国聘网] 成功解析 {len(results)} 条岗位")
    return results


# ============================================================
# Playwright 选择器自动探测
# ============================================================

def _find_job_cards(page, selectors, site_name):
    """
    逐个尝试 CSS 选择器，找到第一个返回结果的即采用。
    如果所有选择器都返回 0 个结果，调用 _auto_detect_job_cards 进行 DOM 自动分析。

    参数:
        page: Playwright page 对象
        selectors: list[str]，CSS 选择器列表（按优先级排序）
        site_name: 站点名称，用于日志

    返回:
        list: 匹配到的 ElementHandle 列表（可能为空）
    """
    for sel in selectors:
        try:
            cards = page.query_selector_all(sel)
            if cards:
                print(f"[{site_name}] 选择器命中: {sel} -> {len(cards)} 个卡片")
                return cards
        except Exception as e:
            print(f"[{site_name}] 选择器异常 ({sel}): {e}")

    print(f"[{site_name}] 所有预设选择器均未匹配，启动 DOM 自动探测...")
    detected = _auto_detect_job_cards(page, site_name)
    if detected:
        print(f"[{site_name}] DOM 自动探测成功，找到 {len(detected)} 个卡片")
    else:
        print(f"[{site_name}] DOM 自动探测也未找到岗位卡片")
    return detected


def _auto_detect_job_cards(page, site_name):
    """
    通过 JavaScript 分析 DOM 结构，自动寻找疑似岗位卡片的元素。

    策略:
    1. 找到所有 class 名包含 job/position/card/item/list 等关键词的元素
    2. 筛选出包含薪资模式（数字+K/元/万）或公司关键词的元素
    3. 取这些元素的共同最近父级容器作为卡片集合
    4. 去重并返回
    """
    try:
        # 用 JS 在页面上执行 DOM 分析
        js_code = """
        () => {
            // Step 1: 收集疑似卡片元素
            const jobKeywords = /job|position|card|item|list|result|vacancy|seek/i;
            const salaryPattern = /\\d+[kK万]|元\\/月|元\\/天|薪|面议|月薪|年薪|\\d+-\\d+/;
            const companyPattern = /公司|集团|有限|科技|股份|控股|研究院|研究所|中心|局|院|厂/;

            // 排除非内容区域（footer、nav、header、sidebar、filter等，含文本匹配）
            const excludeClsPattern = /footer|nav-bar|sidebar|header|popup|modal|copyright|banner|notice-bar|login|register|toolbar|menu|filter-condition|pagination|page/;
            const excludeTextPattern = /版权所有|Copyright|copyight|ICP备|备案号|公司首页|官方微博|官方微信|APP下载|常见问题|清空$|^城市$/;
            function isExcluded(el) {
                // 检查 class 名排除
                let current = el;
                while (current && current !== document.body) {
                    const cls = (current.className || '').toString();
                    if (excludeClsPattern.test(cls)) return true;
                    current = current.parentElement;
                }
                // 检查文本内容排除（避免版权/备案/导航/筛选栏等非岗位区域）
                const text = (el.innerText || '').substring(0, 200);
                if (excludeTextPattern.test(text)) return true;
                return false;
            }

            // 所有包含 job 相关 class 的元素
            const candidates = [];
            const allElements = document.querySelectorAll('*');

            for (const el of allElements) {
                // 先排除非内容区域（在 class 检查之前）
                if (isExcluded(el)) continue;

                const cls = el.className || '';
                const clsStr = typeof cls === 'string' ? cls : '';
                if (!jobKeywords.test(clsStr)) continue;

                // 检查元素内是否包含疑似薪资或公司信息
                const text = el.innerText || '';
                if (text.length < 10 || text.length > 1000) continue;

                const hasSalary = salaryPattern.test(text);
                const hasCompany = companyPattern.test(text);

                if (hasSalary || hasCompany) {
                    candidates.push({
                        el: el,
                        cls: clsStr,
                        text: text.substring(0, 200),
                        depth: el.getBoundingClientRect().width,
                    });
                }
            }

            // Step 2: 如果候选元素太少，用更宽泛的策略 —— 找所有 <li> 或 <div> 中包含薪资信息的
            if (candidates.length < 2) {
                const listItems = document.querySelectorAll('li, div[class]');
                for (const el of listItems) {
                    const text = el.innerText || '';
                    if (text.length < 15 || text.length > 800) continue;
                    if (el.children.length > 10) continue; // 跳过容器级元素

                    const hasSalary = salaryPattern.test(text);
                    if (!hasSalary) continue;

                    // 确保不是嵌套在已选元素里的子元素
                    const cls = el.className || '';
                    const clsStr = typeof cls === 'string' ? cls : '';
                    candidates.push({
                        el: el,
                        cls: clsStr,
                        text: text.substring(0, 200),
                        depth: el.getBoundingClientRect().width,
                    });
                }
            }

            // Step 3: 去重 —— 移除被其他候选元素包含的子元素
            const unique = [];
            for (let i = 0; i < candidates.length; i++) {
                const el1 = candidates[i].el;
                let isParent = false;
                for (let j = 0; j < candidates.length; j++) {
                    if (i === j) continue;
                    const el2 = candidates[j].el;
                    // 如果 el1 包含 el2，则 el1 是父容器，跳过
                    if (el1.contains(el2) && el1 !== el2) {
                        isParent = true;
                        break;
                    }
                }
                if (!isParent) {
                    unique.push(candidates[i]);
                }
            }

            // Step 4: 限制数量
            const result = unique.slice(0, 30);

            // 返回选择器信息供 Playwright 使用
            // 为每个找到的元素标记一个 data 属性
            const markedSelectors = [];
            result.forEach((item, idx) => {
                const marker = `__auto_card_${idx}`;
                item.el.setAttribute('data-auto-card', marker);
                markedSelectors.push({
                    marker: marker,
                    cls: item.cls,
                    preview: item.text.substring(0, 100),
                });
            });

            return markedSelectors;
        }
        """
        markers = page.evaluate(js_code)

        if not markers:
            return []

        # Python 层面的二次过滤：排除疑似非岗位区域（版权/页脚/筛选栏等）
        _EXCLUDE_CLS_RE = re.compile(
            r"footer|nav-bar|sidebar|header|popup|modal|copyright|banner"
            r"|notice-bar|login|register|toolbar|menu|pagination|page"
            r"|filter-condition|bg-|navbar",
            re.IGNORECASE,
        )
        _EXCLUDE_TXT_RE = re.compile(
            r"版权所有|Copyright|copyight|ICP备|备案号|公司首页|官方微博"
            r"|官方微信|APP下载|常见问题|清空\s*$|^城市\s*$"
            r"|请求\s*ID|Security Verification|Captcha|验证码"
            r"|DNS解析失败|Network is unreachable|connection refused",
        )
        # 排除极短/无意义的 class 名（如 't j', 'a b' 等单字母组合）
        _EXCLUDE_CLS_INVALID = re.compile(r"^[a-z]\s[a-z]$", re.IGNORECASE)

        cards = []
        for m in markers:
            cls_str = m.get("cls", "")
            prw_str = m.get("preview", "")

            if _EXCLUDE_CLS_RE.search(cls_str):
                continue
            if _EXCLUDE_CLS_INVALID.match(cls_str):
                continue
            if _EXCLUDE_TXT_RE.search(prw_str):
                continue

            sel = f"[data-auto-card='{m['marker']}']"
            try:
                el = page.query_selector(sel)
                if el:
                    cards.append(el)
                    # 记录探测到的选择器，便于下次更新预设选择器
                    print(f"  [{site_name}] 探测到卡片: class='{cls_str[:60]}' 预览: {prw_str[:50]}...")
            except Exception:
                continue

        return cards

    except Exception as e:
        print(f"  [{site_name}] DOM 自动探测异常: {e}")
        return []


def _extract_text(element, selectors, attr=None):
    """
    从元素中按选择器优先级提取文本或属性。
    参数:
        element: Playwright ElementHandle
        selectors: list[str]，CSS 选择器列表
        attr: 如果指定则提取属性值，否则提取 inner_text
    返回:
        str: 提取到的文本/属性值，失败返回空字符串
    """
    for sel in selectors:
        try:
            child = element.query_selector(sel)
            if child:
                if attr:
                    val = child.get_attribute(attr)
                    if val:
                        return val.strip()
                else:
                    text = child.inner_text()
                    if text.strip():
                        return text.strip()
        except Exception:
            continue
    return ""


# ============================================================
# 爬虫：BOSS直聘（Playwright）
# ============================================================

def _load_saved_cookies(site_name, data_dir=None):
    """
    加载本地保存的已登录 Cookie 文件。
    文件位置: data/{site_name}_cookies.json
    返回 list[dict] 或空列表。
    """
    import os
    if data_dir is None:
        data_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "data")
    cookie_file = os.path.join(data_dir, f"{site_name}_cookies.json")
    if not os.path.exists(cookie_file):
        return []
    try:
        with open(cookie_file, "r", encoding="utf-8") as f:
            cookies = json.load(f)
        if isinstance(cookies, list) and len(cookies) > 0:
            print(f"  [Cookie] 已加载 {site_name} 的 {len(cookies)} 个 cookies ({cookie_file})")
            return cookies
    except Exception as e:
        print(f"  [Cookie] 解析失败 ({cookie_file}): {e}")
    return []


def crawl_boss(keyword, city=None):
    """
    爬取BOSS直聘岗位数据。
    需要使用 Playwright 渲染 JS 页面。
    """
    print(f"\n[BOSS直聘] 开始爬取，关键词: {keyword}，城市: {city}")
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        print("[BOSS直聘] Playwright 未安装，跳过")
        return []

    encoded_keyword = urllib.parse.quote(keyword)
    url = f"https://www.zhipin.com/web/geek/job?query={encoded_keyword}"
    if city:
        # BOSS直聘城市编码映射（部分主要城市）
        city_code_map = {
            "北京": "101010100", "上海": "101020100", "广州": "101280100",
            "深圳": "101280600", "杭州": "101210100", "成都": "101270100",
            "长沙": "101250100", "武汉": "101200100", "南京": "101190100",
            "西安": "101110100", "苏州": "101190400", "重庆": "101040100",
            "天津": "101030100", "郑州": "101180100", "青岛": "101120200",
            "沈阳": "101070100", "大连": "101070200", "厦门": "101230200",
            "合肥": "101220100", "济南": "101120100", "哈尔滨": "101050100",
            "福州": "101230100", "昆明": "101290100", "贵阳": "101260100",
            "南宁": "101300100", "石家庄": "101090100", "太原": "101100100",
            "长春": "101060100", "南昌": "101240100", "兰州": "101160100",
            "海口": "101310100", "呼和浩特": "101080100", "银川": "101170100",
            "西宁": "101150100", "乌鲁木齐": "101130100", "拉萨": "101140100",
        }
        city_code = city_code_map.get(city.replace("市", ""))
        if city_code:
            url += f"&city={city_code}"

    results = []
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(
                headless=True,
                args=[
                    "--no-sandbox",
                    "--disable-setuid-sandbox",
                    "--disable-blink-features=AutomationControlled",
                ],
            )
            context = browser.new_context(
                user_agent=COMMON_UA,
                viewport={"width": 1920, "height": 1080},
                locale="zh-CN",
            )
            page = context.new_page()

            # 尝试加载已保存的 BOSS直聘 Cookie（可选，需手动准备）
            saved_cookies = _load_saved_cookies("boss")
            if saved_cookies:
                context.add_cookies(saved_cookies)
                print("[BOSS直聘] 已注入保存的 cookies，尝试已登录状态搜索")
            else:
                print("[BOSS直聘] 未找到保存的 cookies，将以游客模式访问（大概率被拦截）")

            # 设置反检测
            page.add_init_script("""
                Object.defineProperty(navigator, 'webdriver', { get: () => undefined });
                Object.defineProperty(navigator, 'plugins', { get: () => [1, 2, 3, 4, 5] });
                Object.defineProperty(navigator, 'languages', { get: () => ['zh-CN', 'zh', 'en'] });
            """)

            print(f"[BOSS直聘] 正在加载页面: {url}")
            try:
                page.goto(url, wait_until="domcontentloaded", timeout=SITE_TIMEOUT * 1000)
            except Exception as e:
                print(f"[BOSS直聘] 页面加载超时或失败: {e}")
                browser.close()
                return []

            # 等待页面渲染
            wait_seconds = random.randint(3, 5)
            print(f"[BOSS直聘] 等待页面渲染 {wait_seconds} 秒...")
            page.wait_for_timeout(wait_seconds * 1000)

            # 检查是否被反爬或要求登录
            try:
                page_content = page.content()
            except Exception:
                print("[BOSS直聘] ⚠️ 页面仍在导航中，跳过该网站")
                browser.close()
                return []
            page_title = page.title()
            if "安全验证" in page_content or "验证码" in page_content or "slider" in page_content.lower():
                print("[BOSS直聘] ⚠️ 被反爬拦截，跳过该网站")
                browser.close()
                return []
            if "登录" in page_title or "注册" in page_title:
                print("[BOSS直聘] ⚠️ 需要登录才能查看搜索结果（已重定向到登录页），跳过该网站")
                browser.close()
                return []

            # 尝试获取岗位列表 —— 使用选择器自动探测机制
            # BOSS直聘的岗位卡片选择器（按优先级排序，逐个尝试）
            boss_selectors = [
                ".job-card-wrapper",
                ".search-job-result li",
                ".job-list li",
                "[class*='job-card']",
                "[class*='job-item']",
                ".job-card-body",
                "ul.job-list-box li",
                ".search-job-result .job-card-wrapper",
                "li[ka]",
            ]
            job_cards = _find_job_cards(page, boss_selectors, "BOSS直聘")

            # ★ 批量 JS 提取：避免逐个 ElementHandle 操作时 context 丢失
            # BOSS直聘页面会动态重渲染 DOM，导致 ElementHandle 失效（"Cannot find context"）
            # 改为在单次 page.evaluate 中一次性提取所有卡片数据
            if job_cards:
                print(f"[BOSS直聘] 使用批量JS提取 {min(len(job_cards), 30)} 个卡片数据...")
                try:
                    extracted_data = page.evaluate("""
                        () => {
                            const cards = document.querySelectorAll(
                                'li[ka], .job-card-wrapper, .search-job-result li, ' +
                                '.job-list li, [class*="job-card"], [class*="job-item"], ' +
                                '.job-card-body, ul.job-list-box li'
                            );
                            const results = [];
                            for (const card of cards) {
                                if (results.length >= 30) break;
                                try {
                                    // 岗位名称
                                    let jobTitle = '';
                                    const titleSels = ['.job-name', '.job-title', '[class*="job-name"]', '.job-info .name', 'span[title]'];
                                    for (const sel of titleSels) {
                                        const el = card.querySelector(sel);
                                        if (el && el.innerText.trim()) { jobTitle = el.innerText.trim(); break; }
                                    }

                                    // 公司名称
                                    let companyName = '';
                                    const compSels = ['.company-name', '.company-info', '[class*="company-name"]', '.company-info .name', '[class*="companyName"]'];
                                    for (const sel of compSels) {
                                        const el = card.querySelector(sel);
                                        if (el && el.innerText.trim()) { companyName = el.innerText.trim(); break; }
                                    }

                                    // 薪资
                                    let salary = '';
                                    const salSels = ['.salary', '.job-salary', '[class*="salary"]', '.job-info .salary', '.red'];
                                    for (const sel of salSels) {
                                        const el = card.querySelector(sel);
                                        if (el && el.innerText.trim()) { salary = el.innerText.trim(); break; }
                                    }

                                    // 地点
                                    let workLocation = '';
                                    const areaSels = ['.job-area', '.job-area-wrapper', '[class*="area"]', '.job-info .job-area', '[class*="city"]'];
                                    for (const sel of areaSels) {
                                        const el = card.querySelector(sel);
                                        if (el && el.innerText.trim()) { workLocation = el.innerText.trim(); break; }
                                    }

                                    // 学历
                                    let edu = '';
                                    const eduSels = ['.job-info .edu', '[class*="edu"]', '.job-detail .edu', '[class*="degree"]'];
                                    for (const sel of eduSels) {
                                        const el = card.querySelector(sel);
                                        if (el && el.innerText.trim()) { edu = el.innerText.trim(); break; }
                                    }

                                    // 详情链接
                                    let detailLink = '';
                                    const linkEl = card.querySelector('a.job-card-left, a[href*="job_detail"], a[href*="/job/"], a');
                                    if (linkEl) {
                                        let href = linkEl.getAttribute('href') || '';
                                        if (href && !href.startsWith('http')) {
                                            detailLink = 'https://www.zhipin.com' + href;
                                        } else {
                                            detailLink = href;
                                        }
                                    }

                                    if (jobTitle && companyName) {
                                        results.push({ jobTitle, companyName, salary, workLocation, edu, detailLink });
                                    }
                                } catch(e) { continue; }
                            }
                            return results;
                        }
                    """)

                    print(f"[BOSS直聘] 批量JS提取到 {len(extracted_data)} 条有效数据")

                    for item in extracted_data:
                        try:
                            job_title = item.get("jobTitle", "")
                            company_name = item.get("companyName", "")
                            if not job_title or not company_name:
                                continue

                            salary_range = item.get("salary", "") or "面议"
                            work_location = item.get("workLocation", "") or "全国"
                            edu = item.get("edu", "")
                            detail_link = item.get("detailLink", "")

                            enterprise_type = _determine_enterprise_type(company_name, "BOSS直聘")

                            job_obj = {
                                "job_title": job_title,
                                "company": company_name,
                                "enterprise_type": enterprise_type,
                                "detail_link": detail_link,
                                "salary_range": salary_range,
                                "responsibilities": "详见岗位详情页",
                                "requirements": f"学历要求：{edu}" if edu else "详见岗位详情页",
                                "benefits": "详见岗位详情页",
                                "development": _generate_development(enterprise_type, "BOSS直聘"),
                                "work_location": work_location,
                                "source": "BOSS直聘（真实数据）",
                                "search_keyword": keyword,
                                "_raw_contents": "",
                                "_raw_edu": edu,
                                "_raw_exp": "",
                            }
                            results.append(job_obj)
                        except Exception as e:
                            print(f"  [WARNING] 解析BOSS直聘批量数据失败: {e}")
                            continue

                except Exception as e:
                    print(f"[BOSS直聘] 批量JS提取异常: {e}")
                    # 回退到逐个解析（可能仍会失败，但作为最后手段）
                    print("[BOSS直聘] 回退到逐个ElementHandle解析...")
                    for card in job_cards[:30]:
                        try:
                            job_title = _extract_text(card, [
                                ".job-name", ".job-title", "[class*='job-name']",
                                ".job-info .name", "span[title]",
                            ])
                            company_name = _extract_text(card, [
                                ".company-name", ".company-info", "[class*='company-name']",
                                ".company-info .name", "[class*='companyName']",
                            ])
                            salary_range = _extract_text(card, [
                                ".salary", ".job-salary", "[class*='salary']",
                                ".job-info .salary", ".red",
                            ]) or "面议"
                            work_location = _extract_text(card, [
                                ".job-area", ".job-area-wrapper", "[class*='area']",
                                ".job-info .job-area", "[class*='city']",
                            ])
                            edu = _extract_text(card, [
                                ".job-info .edu", "[class*='edu']", ".job-detail .edu",
                                ".job-info span", "[class*='degree']",
                            ])
                            detail_link = ""
                            link_el = card.query_selector(
                                "a.job-card-left, a[href*='job_detail'], a[href*='/job/'], a"
                            )
                            if link_el:
                                href = link_el.get_attribute("href") or ""
                                if href and not href.startswith("http"):
                                    detail_link = "https://www.zhipin.com" + href
                                else:
                                    detail_link = href

                            if not job_title or not company_name:
                                continue

                            enterprise_type = _determine_enterprise_type(company_name, "BOSS直聘")
                            job_obj = {
                                "job_title": job_title, "company": company_name,
                                "enterprise_type": enterprise_type, "detail_link": detail_link,
                                "salary_range": salary_range, "responsibilities": "详见岗位详情页",
                                "requirements": f"学历要求：{edu}" if edu else "详见岗位详情页",
                                "benefits": "详见岗位详情页",
                                "development": _generate_development(enterprise_type, "BOSS直聘"),
                                "work_location": work_location or "全国",
                                "source": "BOSS直聘（真实数据）", "search_keyword": keyword,
                                "_raw_contents": "", "_raw_edu": edu, "_raw_exp": "",
                            }
                            results.append(job_obj)
                        except Exception as e2:
                            print(f"  [WARNING] 逐个解析也失败: {e2}")
                            continue

            browser.close()
    except Exception as e:
        print(f"[BOSS直聘] 爬取异常: {e}")
        return []

    print(f"[BOSS直聘] 成功解析 {len(results)} 条岗位")
    return results


#!/usr/bin/env python3
# ... license/copyright info ...
# 爬虫：智联招聘（Playwright）
# ============================================================

def crawl_zhaopin(keyword, city=None):
    """
    爬取智联招聘岗位数据。
    DOM 结构（2026-07 验证）：
    .joblist-box__item → 岗位卡片
      .jobinfo__name-row → 岗位名称
      .jobinfo__salary → 薪资
      .jobinfo__other-info-item:nth-child(1) → 地点
      .jobinfo__other-info-item:nth-child(2) → 经验
      .jobinfo__other-info-item:nth-child(3) → 学历
      .companyinfo__name → 公司名称
      .companyinfo__tag → 企业类型标签
    """
    print(f"\n[智联招聘] 开始爬取，关键词: {keyword}，城市: {city}")
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        print("[智联招聘] Playwright 未安装，跳过")
        return []

    encoded_keyword = urllib.parse.quote(keyword)
    url = f"https://sou.zhaopin.com/?kw={encoded_keyword}"
    if city:
        url += f"&jl={urllib.parse.quote(city)}"

    results = []
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(
                headless=True,
                args=[
                    "--no-sandbox",
                    "--disable-setuid-sandbox",
                    "--disable-blink-features=AutomationControlled",
                ],
            )
            context = browser.new_context(
                user_agent=COMMON_UA,
                viewport={"width": 1920, "height": 1080},
                locale="zh-CN",
                timezone_id="Asia/Shanghai",
            )
            page = context.new_page()

            # 反检测
            page.add_init_script("""
                Object.defineProperty(navigator, 'webdriver', { get: () => undefined });
                Object.defineProperty(navigator, 'plugins', { get: () => [1, 2, 3, 4, 5] });
                Object.defineProperty(navigator, 'languages', { get: () => ['zh-CN', 'zh', 'en'] });
                window.chrome = { runtime: {}, loadTimes: function(){}, csi: function(){}, app: {} };
            """)

            # 先访问首页建立 cookie
            try:
                page.goto("https://www.zhaopin.com/", wait_until="domcontentloaded", timeout=15000)
                page.wait_for_timeout(2000)
            except Exception:
                pass

            print(f"[智联招聘] 正在加载搜索页: {url}")
            try:
                page.goto(url, wait_until="domcontentloaded", timeout=SITE_TIMEOUT * 1000)
            except Exception as e:
                print(f"[智联招聘] 搜索页加载超时: {e}")
                browser.close()
                return []

            page.wait_for_timeout(5000)

            page_title = page.title()
            print(f"[智联招聘] 页面标题: {page_title}")

            # 快速检测安全验证页面（不等待 DOM 解析）
            if "Security Verification" in page_title or "安全验证" in page_title:
                print("[智联招聘] ⚠️ 触发安全验证页面（GitHub Actions IP 被限制），跳过")
                browser.close()
                return []

            # 智联招聘岗位卡片 —— 已验证选择器 (2026-07)
            job_cards = page.query_selector_all(".joblist-box__item")
            print(f"[智联招聘] ✓ 找到 {len(job_cards)} 个岗位卡片 (.joblist-box__item)")

            # 如果没有卡片，再检查是否被反爬
            if not job_cards:
                try:
                    page_content = page.content()
                except Exception:
                    print("[智联招聘] ⚠️ 页面仍在加载中，跳过")
                    browser.close()
                    return []
                # 真正的反爬特征（不在正常页面源码中）
                real_block_signs = ["请完成安全验证", "拖动滑块", "geetest_embed", "滑动验证", "Security Verification", "gcaptcha"]
                if any(s in page_content[:10000] for s in real_block_signs):
                    print("[智联招聘] ⚠️ 触发反爬验证，跳过该网站")
                    browser.close()
                    return []
                else:
                    html_snippet = page_content[:2000]
                    print(f"[智联招聘] ⚠️ 未找到卡片（非反爬原因），页面HTML前2000字符: {html_snippet}")

            for card in job_cards[:30]:
                try:
                    # 岗位名称 .jobinfo__name-row
                    job_title = _extract_text(card, [
                        ".jobinfo__name-row", ".jobinfo__name",
                        ".job-name", "[class*='jobName']",
                    ])

                    # 公司名称 .companyinfo__name
                    company_name = _extract_text(card, [
                        ".companyinfo__name", ".company-name",
                        "[class*='companyName']", ".company-info a",
                    ])

                    # 薪资 .jobinfo__salary
                    salary_range = _extract_text(card, [
                        ".jobinfo__salary", ".salary", "[class*='salary']",
                    ]) or "面议"

                    # 地点/经验/学历 —— .jobinfo__other-info-item
                    other_items = card.query_selector_all(".jobinfo__other-info-item")
                    work_location = ""
                    exp_text = ""
                    edu = ""
                    for idx, item in enumerate(other_items[:4]):
                        t = item.inner_text().strip() if item else ""
                        if idx == 0:
                            work_location = t  # 第一个：地点
                        elif idx == 1:
                            exp_text = t        # 第二个：经验
                        elif idx == 2:
                            edu = t             # 第三个：学历

                    # 公司标签（企业类型信息）
                    company_tag = _extract_text(card, [
                        ".companyinfo__tag", "[class*='companyTag']",
                    ])

                    # 详情链接 —— 尝试多种方式
                    detail_link = ""
                    # 方式1: card 本身是 a 标签
                    if card.evaluate("el => el.tagName") == "A":
                        href = card.get_attribute("href") or ""
                        detail_link = href
                    # 方式2: card 内的 a 标签
                    if not detail_link:
                        detail_link = _extract_text(card, [
                            "a[href*='jobs.zhaopin.com']",
                            "a[data-link]", "a.joblist-box__item",
                        ], attr="href")
                    # 方式3: companyinfo__name 的 a 标签
                    if not detail_link:
                        detail_link = _extract_text(card, [
                            ".companyinfo__name",
                        ], attr="href")

                    if not detail_link:
                        detail_link = url  # 兜底

                    if not job_title or not company_name:
                        continue

                    # 根据公司标签或名称判断企业类型
                    enterprise_type = _determine_enterprise_type(
                        company_tag + " " + company_name, "智联招聘"
                    )

                    # 构建 requirements 字符串
                    requirements_parts = []
                    if edu:
                        requirements_parts.append(f"学历要求：{edu}")
                    if exp_text:
                        requirements_parts.append(f"经验要求：{exp_text}")
                    requirements = "\n".join(requirements_parts) if requirements_parts else "详见岗位详情页"

                    job_obj = {
                        "job_title": job_title,
                        "company": company_name,
                        "enterprise_type": enterprise_type,
                        "detail_link": detail_link,
                        "salary_range": salary_range if salary_range else "面议",
                        "responsibilities": "详见岗位详情页",
                        "requirements": requirements,
                        "benefits": "详见岗位详情页",
                        "development": _generate_development(enterprise_type, "智联招聘"),
                        "work_location": work_location or city or "全国",
                        "source": "智联招聘（真实数据）",
                        "search_keyword": keyword,
                        "_raw_contents": company_tag or "",
                        "_raw_edu": edu,
                        "_raw_exp": exp_text,
                    }
                    results.append(job_obj)
                except Exception as e:
                    print(f"  [WARNING] 解析智联招聘岗位卡片失败: {e}")
                    continue

            browser.close()
    except Exception as e:
        print(f"[智联招聘] 爬取异常: {e}")
        return []

    print(f"[智联招聘] ✓ 成功解析 {len(results)} 条岗位")
    return results


# ============================================================
# 爬虫：国资委央企招聘（静态HTML）
# ============================================================

def crawl_sasac(keyword, city=None):
    """
    爬取国资委央企招聘公告。
    静态 HTML，使用 requests + BeautifulSoup。
    """
    print(f"\n[国资委] 开始爬取，关键词: {keyword}")
    # 尝试多个可能的国资委招聘页面URL
    sasac_urls = [
        "http://www.sasac.gov.cn/n2588035/n2588105/index.html",
        "https://www.sasac.gov.cn/n2588035/n2588105/index.html",
        "http://www.sasac.gov.cn/n2588035/c15456054/list.html",
        "http://www.sasac.gov.cn/n2588035/n2588105/c15456054/list.html",
    ]

    resp = None
    base_url = sasac_urls[0]
    for url in sasac_urls:
        resp = _safe_request(url)
        if resp:
            base_url = url
            break
        _random_delay()

    if not resp:
        print("[国资委] 所有URL请求均失败，跳过")
        return []

    try:
        resp.encoding = resp.apparent_encoding or "utf-8"
        soup = BeautifulSoup(resp.text, "lxml")
    except Exception as e:
        print(f"[国资委] HTML解析失败: {e}")
        return []

    results = []

    # 尝试多种选择器匹配公告列表
    items = soup.select("ul.list li, .list li, .news_list li, .content li, li a")
    print(f"[国资委] 找到 {len(items)} 个列表项")

    for item in items[:50]:
        try:
            link_el = item if item.name == "a" else item.find("a")
            if not link_el:
                continue

            title = link_el.get_text(strip=True)
            href = link_el.get("href", "")

            if not title or len(title) < 5:
                continue

            # 关键词过滤
            if keyword and keyword.lower() not in title.lower():
                # 宽松匹配：关键词的每个字都在标题中
                kw_chars = set(keyword.replace(" ", ""))
                title_chars = set(title)
                if len(kw_chars & title_chars) / max(len(kw_chars), 1) < 0.3:
                    continue

            # 补全链接
            if href and not href.startswith("http"):
                if href.startswith("/"):
                    detail_link = "http://www.sasac.gov.cn" + href
                else:
                    detail_link = "http://www.sasac.gov.cn/n2588035/n2588105/" + href
            else:
                detail_link = href

            # 尝试获取发布单位（从父元素或标题中提取）
            parent_text = ""
            if item.parent:
                parent_text = item.parent.get_text(strip=True)
            # 尝试从标题中提取公司名
            company_name = "国资委央企"
            m = re.match(r"^(.+?)(?:招聘|招录|招考|公告|通知)", title)
            if m:
                company_name = m.group(1).strip()

            job_obj = {
                "job_title": title,
                "company": company_name,
                "enterprise_type": "国企",
                "detail_link": detail_link,
                "salary_range": "详见公告",
                "responsibilities": "详见岗位详情页",
                "requirements": "详见岗位详情页",
                "benefits": "央企福利待遇（五险二金/年终奖/带薪年假）",
                "development": "央企平台稳定，晋升通道清晰，福利保障完善",
                "work_location": city or "全国",
                "source": "国资委央企招聘（真实数据）",
                "search_keyword": keyword,
                "_raw_contents": parent_text,
                "_raw_edu": "",
                "_raw_exp": "",
            }
            results.append(job_obj)
        except Exception as e:
            print(f"  [WARNING] 解析国资委岗位失败: {e}")
            continue

    print(f"[国资委] 成功解析 {len(results)} 条岗位")
    return results


# ============================================================
# 爬虫：央企招聘官网（中石油/南方电网/国家电网等）
# ============================================================

def crawl_national_soe(keyword, city=None):
    """
    爬取央企招聘官网（中石油、南方电网、国家电网、中石化等）。
    逐个尝试每个央企的招聘页面，使用 requests + BeautifulSoup 解析。
    如果静态页面解析失败，尝试用 Playwright 渲染。
    """
    print(f"\n[央企官网] 开始爬取 {len(NATIONAL_SOE_SOURCES)} 个央企招聘网站，关键词: {keyword}，城市: {city}")
    results = []

    for soe in NATIONAL_SOE_SOURCES:
        soe_name = soe["name"]
        soe_urls = soe.get("urls", [soe.get("url", "")])
        print(f"  [{soe_name}] 尝试爬取...")
        soe_results = []
        had_http_response = False  # 是否收到过非网络错误的HTTP响应

        for url in soe_urls:
            resp = _safe_request(url, verify_ssl=False)
            if not resp:
                # 检查是否是网络错误（非HTTP错误）—— DNS/连接/超时
                # 如果是网络错误，标记为不可达。如果是HTTP错误(404/412/502)，服务器存在。
                # 我们在 _safe_request 中无法区分，但通过返回值是否为None来统一处理。
                _random_delay()
                continue

            had_http_response = True
            try:
                resp.encoding = resp.apparent_encoding or "utf-8"
                soup = BeautifulSoup(resp.text, "lxml")

                # 通用解析策略：寻找包含"招聘""岗位""职位"的链接和列表项
                # 策略1: 寻找新闻/公告列表
                items = soup.select(
                    "ul li a, .list li a, .news_list li a, "
                    ".content li a, .job-list li, [class*='job'] li, "
                    "[class*='recruit'] li, [class*='position'] li, "
                    "table tr, .article-list li"
                )

                for item in items[:20]:
                    try:
                        link_el = item if item.name == "a" else item.find("a")
                        if not link_el:
                            continue

                        title = link_el.get_text(strip=True)
                        href = link_el.get("href", "")

                        if not title or len(title) < 4:
                            continue

                        # 过滤：包含招聘/岗位/职位/应届/社招等关键词
                        recruit_kw = ["招聘", "岗位", "职位", "应届", "社招", "校园", "招录",
                                      "公告", "招考", "招贤", "招新", "人才", "招工"]
                        if not any(kw in title for kw in recruit_kw):
                            continue

                        # 关键词匹配（宽松）
                        if keyword:
                            kw_lower = keyword.lower()
                            title_lower = title.lower()
                            # 关键词直接匹配 或 关键词分词匹配
                            kw_words = [w for w in re.split(r"[/\-_,，\s]+", keyword) if len(w) >= 2]
                            if kw_lower not in title_lower and not any(w in title_lower for w in kw_words):
                                # 不严格过滤，保留所有招聘信息（因为央企招聘公告标题可能不含具体岗位）
                                pass

                        # 补全链接
                        if href and not href.startswith("http"):
                            if href.startswith("/"):
                                detail_link = url.rstrip("/").rsplit("/", 1)[0].rsplit("/", 1)[0] + href
                            else:
                                detail_link = url.rstrip("/") + "/" + href
                        else:
                            detail_link = href

                        # 获取周边文本作为描述
                        parent_text = ""
                        if item.parent:
                            parent_text = item.parent.get_text(strip=True)[:500]

                        job_obj = {
                            "job_title": title,
                            "company": soe_name,
                            "enterprise_type": "国企",
                            "detail_link": detail_link,
                            "salary_range": "详见公告",
                            "responsibilities": "详见岗位详情页",
                            "requirements": "详见岗位详情页",
                            "benefits": "央企福利待遇（五险二金/年终奖/带薪年假/企业年金）",
                            "development": "央企平台稳定，晋升通道清晰，福利保障完善",
                            "work_location": _detect_city_from_text(title + " " + parent_text) or city or "全国",
                            "source": f"{soe_name}官网（真实数据）",
                            "search_keyword": keyword,
                            "_raw_contents": parent_text,
                            "_raw_edu": "",
                            "_raw_exp": "",
                        }
                        soe_results.append(job_obj)
                    except Exception as e:
                        continue

                if soe_results:
                    print(f"    [{soe_name}] 从 {url} 解析到 {len(soe_results)} 条")
                    break  # 成功获取就不再尝试备用URL

            except Exception as e:
                print(f"    [{soe_name}] 解析失败 ({url}): {e}")

            _random_delay()

        # 如果静态解析失败且有HTTP响应，尝试 Playwright 兜底
        # 如果从未收到过HTTP响应（纯网络错误），跳过 Playwright 避免浪费20秒超时
        if not soe_results and had_http_response:
            soe_results = _crawl_soe_with_playwright(soe_name, soe_urls, keyword, city)
        elif not soe_results and not had_http_response:
            print(f"    [{soe_name}] 所有URL网络不可达，跳过 Playwright 兜底")

        results.extend(soe_results)
        print(f"  [{soe_name}] 共获取 {len(soe_results)} 条")

    print(f"[央企官网] 总计解析 {len(results)} 条岗位")
    return results


def _crawl_soe_with_playwright(soe_name, urls, keyword, city):
    """用 Playwright 渲染央企招聘页面（静态解析失败时的兜底）"""
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        return []

    results = []
    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(
                headless=True,
                args=["--no-sandbox", "--disable-setuid-sandbox",
                      "--disable-blink-features=AutomationControlled"],
            )
            context = browser.new_context(
                user_agent=COMMON_UA, viewport={"width": 1920, "height": 1080}, locale="zh-CN",
            )
            page = context.new_page()
            page.add_init_script("""
                Object.defineProperty(navigator, 'webdriver', { get: () => undefined });
            """)

            for url in urls:
                try:
                    print(f"    [{soe_name}] Playwright 加载: {url}")
                    page.goto(url, wait_until="domcontentloaded", timeout=10000)  # 已失败过，降到10s
                    page.wait_for_timeout(2000)  # 已失败过，降到2s

                    # 通用选择器探测
                    selectors = [
                        "[class*='job'] li", "[class*='recruit'] li", "[class*='position'] li",
                        ".list li", ".news_list li", "ul.list li a",
                        "[class*='job'] a", "[class*='recruit'] a",
                        "table tr", ".article-list li", ".content-list li",
                    ]
                    cards = _find_job_cards(page, selectors, soe_name)

                    for card in cards[:15]:
                        try:
                            text = card.inner_text().strip()
                            link_el = card if card.evaluate("el => el.tagName") == "A" else card.query_selector("a")
                            href = link_el.get_attribute("href") if link_el else ""
                            title = text[:100] if text else ""

                            if not title or len(title) < 4:
                                continue

                            if href and not href.startswith("http"):
                                href = url.rstrip("/").rsplit("/", 1)[0].rsplit("/", 1)[0] + href if href.startswith("/") else url.rstrip("/") + "/" + href

                            results.append({
                                "job_title": title.split("\n")[0][:80],
                                "company": soe_name,
                                "enterprise_type": "国企",
                                "detail_link": href or url,
                                "salary_range": "详见公告",
                                "responsibilities": "详见岗位详情页",
                                "requirements": "详见岗位详情页",
                                "benefits": "央企福利待遇（五险二金/年终奖/带薪年假）",
                                "development": "央企平台稳定，晋升通道清晰，福利保障完善",
                                "work_location": _detect_city_from_text(title + " " + text) or city or "全国",
                                "source": f"{soe_name}官网（真实数据）",
                                "search_keyword": keyword,
                                "_raw_contents": text,
                                "_raw_edu": "",
                                "_raw_exp": "",
                            })
                        except Exception:
                            continue

                    if results:
                        break
                except Exception as e:
                    print(f"    [{soe_name}] Playwright 加载失败 ({url}): {e}")

            browser.close()
    except Exception as e:
        print(f"    [{soe_name}] Playwright 异常: {e}")

    return results


# ============================================================
# 爬虫：地方国企（按用户城市匹配）
# ============================================================

def crawl_local_soe(keyword, city=None):
    """
    根据用户所在城市，爬取该城市的地方国企招聘信息。
    如用户在长沙，爬取长沙银行、湖南银行、湖南建工等。
    """
    if not city:
        print("\n[地方国企] 未提供城市信息，跳过")
        return []

    city_key = city.replace("市", "").strip()
    local_soes = LOCAL_SOE_MAP.get(city_key, [])

    if not local_soes:
        print(f"\n[地方国企] 暂未配置「{city}」的地方国企列表，跳过")
        return []

    print(f"\n[地方国企] 开始爬取「{city}」的 {len(local_soes)} 个地方国企招聘网站")
    results = []

    for soe in local_soes:
        soe_name = soe["name"]
        soe_urls = soe.get("urls", [soe.get("url", "")])

        soe_results = []
        had_http_response = False  # 是否收到过非网络错误的HTTP响应

        # 先直接使用 SOE 配置中的 URL（已经过验证的招聘页面）
        for url in soe_urls:
            if not url:
                continue
            print(f"  [{soe_name}] 爬取 {url}...")
            resp = _safe_request(url, verify_ssl=False)
            if not resp:
                _random_delay()
                continue

            had_http_response = True

            try:
                resp.encoding = resp.apparent_encoding or "utf-8"
                soup = BeautifulSoup(resp.text, "lxml")
                items = soup.select(
                    "ul li a, .list li a, .news_list li a, "
                    "[class*='job'] li, [class*='recruit'] li, "
                    "[class*='position'] li, table tr, "
                    ".article-list li, [class*='career'] li"
                )
                for item in items[:15]:
                    try:
                        link_el = item if item.name == "a" else item.find("a")
                        if not link_el:
                            continue
                        title = link_el.get_text(strip=True)
                        href = link_el.get("href", "")
                        if not title or len(title) < 5:
                            continue
                        # 过滤掉明显非招聘的内容
                        skip_kw = ["招标", "中标", "谈判", "视窗", "loading", "首页", "上一页", "下一页"]
                        if any(kw in title for kw in skip_kw):
                            continue
                        recruit_kw = ["招聘", "岗位", "职位", "应届", "社招", "校园",
                                      "招录", "公告", "招考", "人才", "招工", "简历",
                                      "报名", "应聘", "录用", "录用公示"]
                        if not any(kw in title for kw in recruit_kw):
                            continue
                        if href and not href.startswith("http"):
                            href = url.rstrip("/") + (href if href.startswith("/") else "/" + href)
                        parent_text = item.parent.get_text(strip=True)[:500] if item.parent else ""
                        soe_results.append({
                            "job_title": title, "company": soe_name, "enterprise_type": "国企",
                            "detail_link": href, "salary_range": "详见公告",
                            "responsibilities": "详见岗位详情页", "requirements": "详见岗位详情页",
                            "benefits": "国企福利待遇（五险二金/年终奖/带薪年假）",
                            "development": f"{soe_name}地方国企平台稳定",
                            "work_location": city, "source": f"{soe_name}官网（真实数据）",
                            "search_keyword": keyword, "_raw_contents": parent_text,
                            "_raw_edu": "", "_raw_exp": "",
                        })
                    except Exception:
                        continue

                if soe_results:
                    print(f"    [{soe_name}] 从 {url} 解析到 {len(soe_results)} 条")
                    break

            except Exception as e:
                print(f"    [{soe_name}] 解析失败 ({url}): {e}")

            _random_delay()

        # 如果有HTTP响应但HTML解析失败，用 Playwright 兜底
        # 如果从未收到过HTTP响应（纯网络错误），跳过 Playwright 避免浪费20秒超时
        if not soe_results and had_http_response:
            try:
                from playwright.sync_api import sync_playwright
                with sync_playwright() as p:
                    browser = p.chromium.launch(
                        headless=True,
                        args=["--no-sandbox", "--disable-setuid-sandbox",
                              "--disable-blink-features=AutomationControlled",
                              "--ignore-certificate-errors"],
                    )
                    context = browser.new_context(
                        user_agent=COMMON_UA, viewport={"width": 1920, "height": 1080},
                        locale="zh-CN", ignore_https_errors=True,
                    )
                    context.add_init_script("""
                        Object.defineProperty(navigator, 'webdriver', { get: () => undefined });
                        window.chrome = { runtime: {}, app: {} };
                    """)
                    page = context.new_page()
                    for url in soe_urls:
                        if not url:
                            continue
                        print(f"    [{soe_name}] Playwright 加载 {url}...")
                        try:
                            page.goto(url, wait_until="domcontentloaded", timeout=20000)
                            page.wait_for_timeout(4000)
                            html = page.content()
                            soup = BeautifulSoup(html, "lxml")
                            items = soup.select(
                                "ul li a, .list li a, .news_list li a, "
                                "[class*='job'] li, [class*='recruit'] li, "
                                "[class*='position'] li, a[href]"
                            )
                            for item in items[:12]:
                                try:
                                    link_el = item if item.name == "a" else item.find("a")
                                    if not link_el:
                                        continue
                                    title = link_el.get_text(strip=True)
                                    href = link_el.get("href", "")
                                    if not title or len(title) < 5:
                                        continue
                                    skip_kw = ["招标", "中标", "谈判", "视窗", "loading", "首页", "上一页"]
                                    if any(kw in title for kw in skip_kw):
                                        continue
                                    recruit_kw = ["招聘", "岗位", "职位", "应届", "社招", "校园",
                                                  "招录", "公告", "招考", "人才", "招工",
                                                  "报名", "应聘", "录用公示"]
                                    if not any(kw in title for kw in recruit_kw):
                                        continue
                                    if href and not href.startswith("http"):
                                        href = url.rstrip("/") + (href if href.startswith("/") else "/" + href)
                                    soe_results.append({
                                        "job_title": title, "company": soe_name,
                                        "enterprise_type": "国企", "detail_link": href,
                                        "salary_range": "详见公告", "responsibilities": "详见岗位详情页",
                                        "requirements": "详见岗位详情页",
                                        "benefits": "国企福利待遇（五险二金/年终奖/带薪年假）",
                                        "development": f"{soe_name}地方国企平台稳定",
                                        "work_location": city, "source": f"{soe_name}官网（真实数据）",
                                        "search_keyword": keyword, "_raw_contents": "",
                                        "_raw_edu": "", "_raw_exp": "",
                                    })
                                except Exception:
                                    continue
                            if soe_results:
                                print(f"    [{soe_name}] Playwright 解析到 {len(soe_results)} 条")
                                break
                        except Exception as e:
                            print(f"    [{soe_name}] Playwright 失败: {e}")
                    browser.close()
            except ImportError:
                pass

        results.extend(soe_results)
        print(f"  [{soe_name}] 共获取 {len(soe_results)} 条")

    print(f"[地方国企] 总计解析 {len(results)} 条岗位")
    return results


# ============================================================
# ============================================================
# AI 提取：从微信公众号文章提取结构化岗位（火山方舟 DeepSeek V4）
# ============================================================

_ARK_API_KEY = os.environ.get("ARK_API_KEY", "")
_ARK_ENDPOINT = os.environ.get("ARK_ENDPOINT", "")


def _extract_jobs_with_ai(article_text, user_field, user_city):
    """使用火山方舟 DeepSeek V4 Pro 从公众号文章正文提取岗位。"""
    if not _ARK_API_KEY or not _ARK_ENDPOINT:
        return []

    text_snippet = article_text[:3000]

    prompt = f"""你是招聘信息提取助手。从以下公众号文章提取所有招聘岗位。

要求:
1. 只提取与「{user_field}」相关的岗位
2. 优先提取「{user_city}」的岗位
3. 每个岗位输出: {{"job_title":"","company":"","salary":"","location":""}}
4. 无相关岗位返回 []

只返回 JSON 数组:"""

    try:
        resp = requests.post(
            "https://ark.cn-beijing.volces.com/api/v3/chat/completions",
            headers={"Content-Type": "application/json", "Authorization": f"Bearer {_ARK_API_KEY}"},
            json={"model": _ARK_ENDPOINT, "messages": [
                {"role": "user", "content": prompt + "\n" + text_snippet}
            ], "max_tokens": 1500, "temperature": 0.1},
            timeout=30,
        )
        if resp.status_code != 200:
            return []
        content = resp.json()["choices"][0]["message"]["content"]
        match = re.search(r"\[.*\]", content, re.DOTALL)
        return json.loads(match.group()) if match else []
    except Exception:
        return []


# 爬虫：微信公众号招聘文章（通过搜狗微信搜索）
# ============================================================

# 时效性阈值：只取发布在 960 小时（40天）内的文章
_WECHAT_MAX_HOURS = 960

# 已验证的公众号白名单缓存（蓝V认证），避免对同一账号重复检测
_VERIFIED_ACCOUNTS_CACHE = set()
_UNVERIFIED_ACCOUNTS_CACHE = set()


def _parse_sogou_publish_time(time_text):
    """解析搜狗微信搜索结果的发布时间，返回 datetime 或 None。

    搜狗页面上的时间格式：
    - "7月15日" → 今年7月15日
    - "3天前" → 3天前
    - "5小时前" → 5小时前
    - "30分钟前" → 30分钟前
    - "昨天" → 昨天
    """
    if not time_text:
        return None
    text = time_text.strip()

    now = datetime.now()
    try:
        # "X天前"
        m = re.match(r"(\d+)\s*天前", text)
        if m:
            days = int(m.group(1))
            from datetime import timedelta
            return now - timedelta(days=days)

        # "X小时前"
        m = re.match(r"(\d+)\s*小时前", text)
        if m:
            hours = int(m.group(1))
            from datetime import timedelta
            return now - timedelta(hours=hours)

        # "X分钟前"
        m = re.match(r"(\d+)\s*分钟前", text)
        if m:
            minutes = int(m.group(1))
            from datetime import timedelta
            return now - timedelta(minutes=minutes)

        # "昨天"
        if "昨天" in text:
            from datetime import timedelta
            return now - timedelta(days=1)

        # "X月X日"
        m = re.match(r"(\d{1,2})\s*月\s*(\d{1,2})\s*日", text)
        if m:
            month, day = int(m.group(1)), int(m.group(2))
            year = now.year
            # 如果月份比当前月份大，说明是去年的
            if month > now.month:
                year -= 1
            return datetime(year, month, day)

        # "X年X月X日"
        m = re.match(r"(\d{4})\s*年\s*(\d{1,2})\s*月\s*(\d{1,2})\s*日", text)
        if m:
            return datetime(int(m.group(1)), int(m.group(2)), int(m.group(3)))
    except Exception:
        pass

    return None


def _is_within_time_window(publish_time, max_hours=_WECHAT_MAX_HOURS):
    """判断发布时间是否在时间窗口内。"""
    if publish_time is None:
        return False  # 无法解析时间的，保守跳过
    from datetime import timedelta
    cutoff = datetime.now() - timedelta(hours=max_hours)
    return publish_time >= cutoff


def _check_wechat_account_verified(page, article_url):
    """基于关键词分类法判定公众号是否为官方认证账号（官号）。

    策略（按优先级）：
    1. 页面DOM认证标识检测（蓝V图标）
    2. 关键词白名单匹配 → 判定为官号
    3. 关键词黑名单匹配 → 判定为中介号
    4. 综合判断 → 未知则保守拒绝

    返回 (verified: bool, account_name: str, account_type: str)
    account_type: "government"|"enterprise"|"media"|"intermediary"|"unverified"
    """

    # ========== 白名单关键词（匹配到任一即判定为官号） ==========
    _SOE_PREFIXES = ["中国", "中华", "国家", "中央"]
    _SOE_SUFFIXES = ["集团有限公司", "控股集团", "股份有限公司", "总公司"]
    _SOE_INDUSTRY = ["石油", "电网", "铁路", "建筑", "烟草", "航天", "兵器", "船舶", "电力", "电信", "国资"]

    _LOCAL_PREFIXES = ["湖南", "长沙", "湘江", "星城", "湘"]
    _LOCAL_SUFFIXES = ["建设集团", "投资集团", "城发", "城投", "国资", "控股", "发展集团"]

    _GOVERNMENT_KW = ["国有资产监督管理委员会", "人民政府", "人社局", "人力资源社会保障"]

    # 高可信官方关键词（直接判定为官号，不依赖前缀/后缀组合）
    _HARD_OFFICIAL_KW = [
        "南方电网", "国家电网", "国家电投", "中国石油", "中国石化", "中国海油",
        "中国移动", "中国联通", "中国电信", "中国建筑", "中国中铁",
        "中国铁建", "中粮集团", "华润集团", "招商局", "中国航天",
        "中国船舶", "中国兵器", "中国航天科工", "中国电子",
        "中金岭南", "中联重科", "长沙银行", "湖南银行",
        "北京大学", "清华大学",
    ]

    # ========== 黑名单关键词 ==========
    _HARD_INTERMEDIARY_KW = [
        # 教育/培训类
        "教育", "教育科技", "人力资源", "劳务派遣", "人才服务", "求职咨询",
        # 已知公考培训机构
        "中公", "华图", "粉笔", "腰果", "公考",
        # 招聘平台类
        "招聘网", "人才网",
    ]

    # 需结合其他关键词综合判断（单独命中不触发黑名单）
    _SOFT_INTERMEDIARY_KW = [
        "文化传媒", "信息科技", "网络科技",
    ]

    def _count_blacklist_hits(name):
        """统计账号名命中黑名单关键词的次数。"""
        count = 0
        for kw in _HARD_INTERMEDIARY_KW:
            if kw in name:
                count += 1
        for kw in _SOFT_INTERMEDIARY_KW:
            if kw in name:
                count += 1
        return count

    def _has_hard_blacklist(name):
        """检查是否命中硬黑名单（无需组合判断即可判定为中介）。"""
        for kw in _HARD_INTERMEDIARY_KW:
            if kw in name:
                return True
        return False

    def _has_soft_blacklist(name):
        """检查是否命中软黑名单。"""
        for kw in _SOFT_INTERMEDIARY_KW:
            if kw in name:
                return True
        return False

    def _is_official_by_keywords(name):
        """基于关键词判断是否为官方账号。"""
        if not name:
            return False, "unknown"

        # === 高可信官方关键词（直接匹配） ===
        for kw in _HARD_OFFICIAL_KW:
            if kw in name:
                return True, "enterprise"

        # === 央企命名模式：前缀 + (行业词 或 后缀) ===
        has_soe_prefix = any(name.startswith(p) for p in _SOE_PREFIXES)
        has_industry = any(kw in name for kw in _SOE_INDUSTRY)
        has_soe_suffix = any(name.endswith(s) for s in _SOE_SUFFIXES)
        if has_soe_prefix and (has_industry or has_soe_suffix):
            return True, "enterprise"

        # === 地方国企模式：地方前缀 + 国企后缀 ===
        has_local_prefix = any(kw in name for kw in _LOCAL_PREFIXES)
        has_local_suffix = any(kw in name for kw in _LOCAL_SUFFIXES)
        if has_local_prefix and has_local_suffix:
            return True, "enterprise"

        # === 政府/事业单位 ===
        for kw in _GOVERNMENT_KW:
            if kw in name:
                return True, "government"

        return False, "unknown"

    def _is_intermediary_by_keywords(name):
        """基于关键词判断是否为中介/培训类账号。"""
        if not name:
            return False, "unknown"

        # 硬黑名单：直接判定为中介
        if _has_hard_blacklist(name):
            return True, "intermediary"

        # 软黑名单：需要至少命中一个其他黑名单关键词
        if _has_soft_blacklist(name):
            hit_count = _count_blacklist_hits(name)
            if hit_count >= 2:
                # 软黑名单 + 任一其他黑名单关键词 = 中介
                return True, "intermediary"
            # 仅命中软黑名单，不判定

        return False, "unknown"

    # ====== 主逻辑 ======
    account_name = ""
    account_type = "unknown"

    try:
        # 从页面提取公众号昵称
        nickname_el = page.query_selector(".rich_media_meta_nickname, #js_name, .wx_follow_nickname")
        if nickname_el:
            account_name = nickname_el.inner_text().strip()

        # === 第1层：页面DOM认证标识检测（最高优先） ===
        verify_selectors = [
            "#js_profile_qrcode .profile_verify_icon",
            ".rich_media_meta_nickname .verify_icon",
            "#js_verify",
            "[class*='wx_verify']",
        ]
        for sel in verify_selectors:
            try:
                el = page.query_selector(sel)
                if el:
                    title_attr = el.get_attribute("title") or el.get_attribute("aria-label") or ""
                    title_lower = title_attr.lower()
                    if any(kw in title_lower for kw in ["政府", "事业单位", "机关"]):
                        return True, account_name, "government"
                    if any(kw in title_lower for kw in ["企业", "公司", "集团"]):
                        return True, account_name, "enterprise"
                    if any(kw in title_lower for kw in ["媒体", "新闻"]):
                        return True, account_name, "media"
                    # 有认证标识但无法确定类型 → 仍视为企业认证
                    if title_attr:
                        return True, account_name, "enterprise"
                    # title 为空说明是误匹配，继续尝试下一个选择器
            except Exception:
                continue

        # === 第2层：关键词白名单匹配 ===
        is_official, official_type = _is_official_by_keywords(account_name)
        if is_official:
            return True, account_name, official_type

        # === 第3层：关键词黑名单匹配 ===
        is_intermediary, intermediary_type = _is_intermediary_by_keywords(account_name)
        if is_intermediary:
            return False, account_name, intermediary_type

        # === 第4层：保守策略 ===
        # 无法判断 → 保守拒绝（避免中介号漏网）
        if account_name and len(account_name) >= 4:
            # 有一定长度但无法识别 → 默认为未认证
            return False, account_name, "unverified"
        elif account_name:
            # 短名称、无法判断 → 拒绝
            return False, account_name, "suspicious_short"

        return False, account_name, "unknown"
    except Exception:
        return False, account_name, "unknown"
    except Exception:
        return False, account_name, "unknown"


def crawl_wechat_sogou(keyword, city=None):
    """
    ★ 通过 Sogou 普通网页搜索发现微信公众号招聘文章。
    
    ★★★ 搜狗微信搜索已被验证码反爬封锁（2026-07-19 起生效），改用 Sogou 普通网页搜索 + mp.weixin.qq.com 链接检测。

    搜索策略：
    1. 央企 + 城市：中石油长沙招聘、南方电网长沙招聘...
    2. 地方国企 + 城市：长沙银行招聘、湖南银行招聘...
    3. 通用：长沙 国企招聘、Python开发 招聘 长沙
    4. 每个词取 mp.weixin.qq.com 链接，去重后返回
    5. 使用 Playwright 渲染避开搜狗反爬
    """
    print(f"\n[微信公众号] ★ Sogou网页搜索（微信文章链接检测），关键词: {keyword}，城市: {city}")
    print(f"  ★ 搜狗微信搜索已被验证码封锁，改用普通搜索引擎发现微信文章")

    city_key = (city or "").replace("市", "").strip()
    local_soenames = [s["name"] for s in LOCAL_SOE_MAP.get(city_key, [])]

    # 构建搜索词组合 —— 优先搜专业+城市，再搜央企通用
    search_terms = []

    # ★ 1. 关键词 + 城市组合（最高优先，最精准）
    if keyword and city:
        search_terms.append(f"{city} {keyword} 招聘")
        search_terms.append(f"{keyword} {city} 招聘")
        search_terms.append(f"{city} {keyword}")
        search_terms.append(f"{keyword} 招聘 {city}")
    elif keyword:
        search_terms.append(f"{keyword} 招聘")

    # 2. 城市 + 国企通用搜索
    if city:
        search_terms.append(f"{city} 国企 招聘")
        search_terms.append(f"{city} 央企 招聘")
        search_terms.append(f"{city} 事业单位 招聘")
        search_terms.append(f"{city} 招聘公告")
        search_terms.append(f"{city} 招聘信息 2026")

    # 3. 每个央企 + 城市组合
    for soe in NATIONAL_SOE_SOURCES:
        if city:
            search_terms.append(f"{soe['name']} {city} 招聘")
        search_terms.append(f"{soe['name']} 招聘公告")

    # 4. 每个地方国企
    for name in local_soenames:
        search_terms.append(f"{name} 招聘公告")
        if city:
            search_terms.append(f"{name} {city} 招聘")

    # 去重并限制（最多30条搜索）
    search_terms = list(dict.fromkeys(search_terms))[:30]

    results = []
    seen_titles = set()
    _STOP_AFTER_N_RESULTS = 40  # 累计40篇就停

    # 使用 Playwright 渲染搜狗普通网页搜索（比 requests 能绕过基础反爬）
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        print("  [微信公众号] Playwright 未安装，跳过微信文章搜索")
        return []

    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(
                headless=True,
                args=["--no-sandbox", "--disable-setuid-sandbox",
                      "--disable-blink-features=AutomationControlled",
                      "--ignore-certificate-errors"],
            )
            context = browser.new_context(
                user_agent=COMMON_UA, viewport={"width": 1920, "height": 1080},
                locale="zh-CN", ignore_https_errors=True,
            )
            context.add_init_script("""
                Object.defineProperty(navigator, 'webdriver', { get: () => undefined });
                window.chrome = { runtime: {}, app: {} };
            """)
            page = context.new_page()

            for term in search_terms:
                if len(seen_titles) >= _STOP_AFTER_N_RESULTS:
                    print(f"  [微信搜索] 已收集 {len(seen_titles)} 篇，停止搜索")
                    break

                encoded_term = urllib.parse.quote(term)
                # ★ 使用 Sogou 普通网页搜索，而非微信专用搜索
                search_url = f"https://www.sogou.com/web?query={encoded_term}"

                try:
                    page.goto(search_url, wait_until="domcontentloaded", timeout=15000)
                    page.wait_for_timeout(2000)  # 等 JS 渲染完成

                    html = page.content()
                    soup = BeautifulSoup(html, "lxml")

                    # ★ 查找所有 mp.weixin.qq.com 链接
                    wechat_links = []
                    all_as = soup.select("a[href]")
                    for a in all_as:
                        href = a.get("href", "")
                        if "mp.weixin.qq.com" in href:
                            link_text = a.get_text(strip=True)
                            # 过滤空文本和短文本
                            if link_text and len(link_text) >= 4:
                                wechat_links.append({
                                    "title": link_text,
                                    "href": href,
                                })

                    # 也尝试在 parent/sibling 元素中提取更多上下文
                    print(f"  [微信搜索] 「{term}」→ 发现 {len(wechat_links)} 个微信文章链接")

                    # 如果链接文本不够好，尝试从附近元素提取标题
                    if len(wechat_links) <= 2:
                        # 尝试更宽泛的搜索
                        wechat_links_extra = []
                        for a in all_as:
                            href = a.get("href", "")
                            if "mp.weixin.qq.com" in href:
                                link_text = a.get_text(strip=True)
                                if not link_text or len(link_text) < 4:
                                    # 从父元素获取文本
                                    parent = a.parent
                                    if parent:
                                        parent_text = parent.get_text(strip=True)[:100]
                                        if parent_text and len(parent_text) >= 4:
                                            link_text = parent_text
                                if link_text and len(link_text) >= 4:
                                    wechat_links_extra.append({
                                        "title": link_text,
                                        "href": href,
                                    })
                        if wechat_links_extra:
                            wechat_links = wechat_links_extra

                    found = 0
                    for wl in wechat_links:
                        if found >= 5:  # 每个搜索词最多取5篇（严格过滤后实际招聘文章有限）
                            break

                        title = wl["title"]
                        href = wl["href"]

                        if not title or len(title) < 5:
                            continue

                        # 去重（基于标题前50字符）
                        title_key = title[:50]
                        if title_key in seen_titles:
                            continue
                        seen_titles.add(title_key)

                        # ★★★ 严格过滤策略 ★★★
                        # 第1层：先排除明显不相关内容
                        _skip_prefix = ["广告", "推广", "充值", "会员", "下载",
                                        "app", "APP", "二维码", "关注公众号",
                                        "扫码", "小程序", "学完", "速看", "速存",
                                        "收藏!", "收藏！", "建议收藏", "建议收藏！"]
                        if any(title.startswith(kw) for kw in _skip_prefix):
                            continue

                        # 第2层：排除公考培训/备考/经验分享类
                        _exam_kw = ["考试培训", "公考培训", "考公培训",
                                    "辅导班", "培训班", "备考指南", "备考攻略",
                                    "备考经验", "公考之路", "公考经验",
                                    "公务员考试", "公考备考", "考公备考",
                                    "面试经验", "面试技巧", "面试方法",
                                    "面试备考", "面试真题", "笔试真题",
                                    "历年真题", "真题解析", "真题汇总",
                                    "高分技巧", "高分经验", "上岸经验",
                                    "上岸分享", "上岸攻略", "成公经验",
                                    "考公经验", "考编经验", "行测",
                                    "申论", "时政热点", "时政汇总",
                                    "公考干货", "考公干货", "考公攻略",
                                    "如何备考", "怎么备考", "怎么复习",
                                    "复习经验", "复习攻略", "复习方法",
                                    "学习方法", "学习计划", "备考计划",
                                    "考公人", "考公党", "公考人", "公考党",
                                    "在职牛马", "在职考生", "在职备考",
                                    "上岸考生", "上岸学长", "上岸学姐",
                                    "已上岸", "成功上岸", "一次上岸",
                                    "笔记分享", "资料分享", "资源分享",
                                    "资料汇总", "资料合集", "备考资料",
                                    "考公资料", "公考资料", "免费领取",
                                    "限时领取", "免费分享", "速领",
                                    # 培训机构/自媒体
                                    "中公教育", "华图教育", "粉笔公考",
                                    "粉笔教育", "腰果公考", "步知公考",
                                    "半月谈", "超格公考", "青京公考",
                                    "得政", "湘麓法源", "娄上双星",
                                    "娄底中公", "湖南中公", "长沙中公"]
                        if any(kw in title for kw in _exam_kw):
                            continue

                        # 第3层：必须包含真实招聘关键词
                        _hard_recruit_kw = ["招聘", "招录", "招考公告", "招聘公告",
                                            "人才引进", "招才引智", "校园招聘",
                                            "社会招聘", "春季招聘", "秋季招聘",
                                            "招聘简章", "招聘计划", "招聘岗位",
                                            "岗位需求", "招贤纳士", "诚聘",
                                            "招新", "纳新", "报名公告",
                                            "录用公示", "录用名单", "拟录用",
                                            "录用公告", "录用人员"]
                        if not any(kw in title for kw in _hard_recruit_kw):
                            # 放宽：如果包含"国企"+"公告"也算
                            if not ("国企" in title and "公告" in title):
                                continue

                        # 第4层：排除 Sogou 搜索结果元数据（非真实文章标题）
                        _meta_patterns = ["·微信公众号", "·公众号", "投诉举报",
                                          "用户协议", "隐私政策", "服务协议",
                                          "Copyright", "©", "ICP"]
                        if any(kw in title for kw in _meta_patterns):
                            continue

                        # 尝试从标题中提取公司名
                        company_name = "国企公众号"
                        for soe in NATIONAL_SOE_SOURCES:
                            if soe["name"] in title:
                                company_name = soe["name"]
                                break
                        if company_name == "国企公众号":
                            for local_name in local_soenames:
                                if local_name in title:
                                    company_name = local_name
                                    break

                        # 检测城市
                        detected_city = _detect_city_from_text(title)
                        work_location = detected_city if detected_city else (city or "全国")

                        # 从URL中提取可能的账号名
                        account_name = ""
                        try:
                            # mp.weixin.qq.com/s?src=11&timestamp=... 这种格式没有账号名
                            # 实际的公众号文章URL是 mp.weixin.qq.com/s?__biz=... 格式
                            account_name = title.split("·")[0].strip() if "·" in title else ""
                        except Exception:
                            pass

                        # 确保 href 是完整URL
                        if href and not href.startswith("http"):
                            href = "https://" + href.lstrip("/")

                        job_obj = {
                            "job_title": title[:80],
                            "company": company_name,
                            "enterprise_type": "国企",
                            "detail_link": href or search_url,
                            "salary_range": "详见公告",
                            "responsibilities": "详见岗位详情页",
                            "requirements": "详见岗位详情页",
                            "benefits": "国企福利待遇（五险二金/年终奖/带薪年假）",
                            "development": "国企平台稳定，晋升通道清晰",
                            "work_location": work_location,
                            "source": f"微信公众号-Sogou搜索（真实数据）",
                            "search_keyword": keyword,
                            "_raw_contents": "",
                            "_raw_edu": "",
                            "_raw_exp": "",
                            "_publish_time_str": "",
                            "_account_name": account_name or title[:20],
                        }
                        results.append(job_obj)
                        found += 1

                    if found > 0:
                        print(f"    → 收录 {found} 篇 (累计 {len(seen_titles)} 篇)")

                except Exception as e:
                    print(f"    [微信搜索] 「{term}」导航失败: {e}")

                _random_delay()

            browser.close()

    except Exception as e:
        print(f"  [微信公众号] Playwright 搜索失败: {e}")
        return []

    # 尝试用 Playwright 获取文章正文 + AI 提取结构化岗位信息
    if results and _ARK_API_KEY and _ARK_ENDPOINT:
        ai_jobs = _extract_from_wechat_articles_with_ai(results, keyword, city)
        if ai_jobs:
            print(f"[微信公众号] ★ AI 从文章正文提取到 {len(ai_jobs)} 个结构化岗位")
            results = ai_jobs
        else:
            print(f"[微信公众号] ★ AI 未提取到结构化岗位，保留搜索结果")

    print(f"[微信公众号] ★ 总计 {len(results)} 篇招聘文章（来自 {len(search_terms)} 个搜索词）")
    return results


def _extract_from_wechat_articles_with_ai(articles, keyword, city):
    """用 Playwright 获取文章正文 + AI 提取结构化岗位。"""
    ai_jobs = []
    seen = set()

    # Step 1: 用 Playwright 批量获取文章正文和真实链接
    article_texts = _fetch_wechat_articles_with_playwright(articles[:10])

    # Step 2: 分批调用 AI
    # 每批 3 篇，避免文本过长导致 AI 超时
    for i in range(0, len(article_texts), 3):
        chunk = article_texts[i:i+3]
        combined = "\n---\n".join(t["text"] for t in chunk)

        extracted = _extract_jobs_with_ai(combined, keyword, city)
        for ej in extracted:
            job_title = ej.get("job_title", "")
            company = ej.get("company", "国企")
            salary = ej.get("salary", "详见公告")
            location = ej.get("location", city or "全国")

            if not job_title or len(job_title) < 3:
                continue

            # 尝试匹配到原始文章的真实链接
            detail_link = ""
            for at in chunk:
                if job_title in at["text"] or company in at["text"]:
                    detail_link = at.get("wechat_url", "")
                    break

            # 如果没有真实链接，生成搜索链接
            if not detail_link:
                detail_link = _build_search_link(job_title, company, city)

            ai_jobs.append({
                "job_title": job_title[:80],
                "company": company,
                "enterprise_type": "国企",
                "detail_link": detail_link,
                "salary_range": salary,
                "responsibilities": "详见岗位详情页",
                "requirements": "详见岗位详情页",
                "benefits": "国企福利待遇（五险二金/年终奖/带薪年假）",
                "development": "国企平台稳定",
                "work_location": location,
                "source": "微信公众号-AI提取（真实数据）",
                "search_keyword": keyword,
                "_raw_contents": combined[:500],
                "_raw_edu": "",
                "_raw_exp": "",
            })

    return ai_jobs


def _fetch_wechat_articles_with_playwright(articles):
    """用 Playwright 批量打开搜狗微信链接，获取文章正文和真实 mp.weixin.qq.com 链接。"""
    results = []
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        # 降级：只用标题和摘要
        for a in articles:
            title = a.get("job_title", "")
            summary = a.get("_raw_contents", "")
            if title and len(title) > 2:
                results.append({"text": f"标题:{title}\n摘要:{summary}", "wechat_url": ""})
        return results

    try:
        with sync_playwright() as p:
            browser = p.chromium.launch(
                headless=True,
                args=["--no-sandbox", "--disable-setuid-sandbox",
                      "--disable-blink-features=AutomationControlled"],
            )
            context = browser.new_context(
                user_agent=COMMON_UA, viewport={"width": 1920, "height": 1080}, locale="zh-CN",
            )
            page = context.new_page()
            page.add_init_script(
                "Object.defineProperty(navigator, 'webdriver', { get: () => undefined });"
            )

            for article in articles:
                url = article.get("detail_link", "")
                title = article.get("job_title", "")
                summary = article.get("_raw_contents", "")
                account_name = article.get("_account_name", "")

                # ★ 账号白名单快速通道：已被验证过的账号直接通过
                if account_name and account_name in _VERIFIED_ACCOUNTS_CACHE:
                    pass  # 已验证，直接走后续流程
                elif account_name and account_name in _UNVERIFIED_ACCOUNTS_CACHE:
                    print(f"    [资质过滤] 跳过未认证账号: {account_name}")
                    continue

                if not url or "sogou.com" not in url:
                    results.append({"text": f"标题:{title}\n摘要:{summary}", "wechat_url": url if url else ""})
                    continue

                try:
                    # ★ 关键技巧：每篇文章先回搜狗搜索页建立 cookie 和 referer，
                    # 让搜狗以为是真人浏览行为，再跳转文章链接就不会触发反爬
                    page.goto("https://weixin.sogou.com/weixin?type=2", wait_until="domcontentloaded", timeout=10000)
                    page.wait_for_timeout(500)
                    page.goto(url, wait_until="domcontentloaded", timeout=15000)
                    page.wait_for_timeout(3000)
                    final_url = page.url

                    if "mp.weixin.qq.com" in final_url:
                        # 先获取页面HTML，用于时间 + 资质双重校验
                        html = page.content()
                        from bs4 import BeautifulSoup as _BS
                        soup = _BS(html, "lxml")

                        # ★ 时效性校验：从公众号文章页提取发布时间
                        publish_time_selectors = [
                            "#publish_time", ".rich_media_meta_text",
                            "#meta_content time", "[class*='publish']",
                            "time", "em#publish_time",
                        ]
                        article_publish_dt = None
                        for pts in publish_time_selectors:
                            pt_el = soup.select_one(pts)
                            if pt_el:
                                pt_text = pt_el.get_text(strip=True) or pt_el.get("datetime", "")
                                article_publish_dt = _parse_sogou_publish_time(pt_text)
                                if article_publish_dt:
                                    break
                        if not _is_within_time_window(article_publish_dt):
                            ts = article_publish_dt.strftime("%Y-%m-%d") if article_publish_dt else "未知"
                            print(f"    [时效过滤] 跳过: {title[:25]} (发布于 {ts}, 超出{_WECHAT_MAX_HOURS}h)")
                            continue

                        # ★ 账号资质校验：检查是否为蓝V认证官方公众号
                        if account_name and account_name not in _VERIFIED_ACCOUNTS_CACHE:
                            verified, _, acct_type = _check_wechat_account_verified(page, final_url)
                            if verified:
                                _VERIFIED_ACCOUNTS_CACHE.add(account_name)
                                print(f"    [资质校验] ✅ 认证账号: {account_name} ({acct_type})")
                            else:
                                _UNVERIFIED_ACCOUNTS_CACHE.add(account_name)
                                print(f"    [资质过滤] ❌ 未认证账号: {account_name}, 跳过")
                                continue

                        body = soup.find("div", id="js_content")
                        text = body.get_text(separator="\n", strip=True)[:2000] if body else title + "\n" + summary
                        results.append({"text": text, "wechat_url": final_url})
                        print(f"    [微信正文] ✅ {title[:25]}")
                    elif "antispider" in final_url:
                        # 被拦，降级为标题摘要
                        results.append({"text": f"标题:{title}\n摘要:{summary}", "wechat_url": ""})
                        print(f"    [微信正文] ⚠️ 反爬: {title[:20]}")
                    else:
                        results.append({"text": f"标题:{title}\n摘要:{summary}", "wechat_url": ""})
                        print(f"    [微信正文] ⚠️ 未到公众号: {title[:20]}")

                except Exception as e:
                    print(f"    [微信正文] ❌ {title[:20]}: {e}")
                    results.append({"text": f"标题:{title}\n摘要:{summary}", "wechat_url": ""})

            browser.close()

    except Exception as e:
        print(f"    [微信正文] Playwright 整体失败: {e}")
        for a in articles:
            title = a.get("job_title", "")
            summary = a.get("_raw_contents", "")
            if title and len(title) > 2:
                results.append({"text": f"标题:{title}\n摘要:{summary}", "wechat_url": ""})

    return results


def _build_search_link(job_title, company, city):
    """为没有详情链接的岗位生成搜索链接。"""
    query = f"{company} {job_title} {city or ''}".strip()
    encoded = urllib.parse.quote(query)
    return f"https://www.iguopin.com/search?keyword={encoded}"


# ============================================================
# 爬虫：国聘网城市定向搜索
# ============================================================

def crawl_iguopin_by_city(keyword, city):
    """
    国聘网城市定向搜索：使用城市筛选参数精准搜索目标城市岗位。
    比通用搜索更精准，专门用于确保有足够的目标城市岗位。
    """
    if not city:
        return crawl_iguopin(keyword, city)

    print(f"\n[国聘网-城市定向] 搜索「{city}」的「{keyword}」岗位")
    api_url = "https://gp-api.iguopin.com/api/jobs/v1/list"
    headers = {
        "Content-Type": "application/json;charset=UTF-8",
        "Device": "pc",
        "Subsite": "cujiuye",
        "Version": "5.0.0",
        "User-Agent": COMMON_UA,
        "Referer": "https://www.iguopin.com/",
        "Origin": "https://www.iguopin.com",
    }

    # 国聘网城市编码映射
    iguopin_city_map = {
        "北京": "110100", "上海": "310100", "广州": "440100",
        "深圳": "440300", "杭州": "330100", "成都": "510100",
        "长沙": "430100", "武汉": "420100", "南京": "320100",
        "西安": "610100", "苏州": "320500", "重庆": "500100",
        "天津": "120100", "郑州": "410100", "青岛": "370200",
        "沈阳": "210100", "大连": "210200", "厦门": "350200",
        "合肥": "340100", "济南": "370100", "哈尔滨": "230100",
        "福州": "350100", "昆明": "530100", "贵阳": "520100",
        "南宁": "450100", "石家庄": "130100", "太原": "140100",
        "长春": "220100", "南昌": "360100", "兰州": "620100",
        "海口": "460100", "呼和浩特": "150100", "银川": "640100",
        "西宁": "630100", "乌鲁木齐": "650100", "拉萨": "540100",
    }
    city_code = iguopin_city_map.get(city.replace("市", ""))

    # 搜索策略1: 带城市编码的精准搜索 + 城市名关键词补充
    payloads = []
    if city_code:
        payloads.append({"page": 1, "page_size": 50, "keyword": keyword, "city": city_code})
    payloads.append({"page": 1, "page_size": 50, "keyword": f"{keyword} {city}"})

    results = []
    for payload in payloads:
        resp = _safe_request(api_url, method="POST", headers=headers, json_body=payload)
        if not resp:
            continue

        try:
            data = resp.json()
            job_list = data.get("data", {}).get("list", []) if isinstance(data, dict) else []

            for item in job_list:
                try:
                    job_id = item.get("job_id") or ""
                    job_name = item.get("job_name") or ""
                    company_name = item.get("company_name") or ""
                    if not job_name or not company_name:
                        continue

                    # 薪资
                    min_wage = item.get("min_wage") or ""
                    max_wage = item.get("max_wage") or ""
                    if min_wage and max_wage:
                        salary_range = f"{min_wage}-{max_wage}元/月"
                    elif min_wage:
                        salary_range = f"{min_wage}元/月起"
                    else:
                        salary_range = item.get("salary") or "面议"

                    edu = item.get("education_cn") or ""
                    exp = item.get("experience_cn") or ""
                    contents = item.get("contents") or ""

                    district_list = item.get("district_list") or []
                    if isinstance(district_list, list) and district_list:
                        work_location = district_list[0].get("area_cn") or city
                    else:
                        work_location = city  # 城市定向搜索默认为该城市

                    company_info = item.get("company_info") or {}
                    nature = company_info.get("nature_cn", "") if isinstance(company_info, dict) else ""

                    enterprise_type = _determine_enterprise_type(nature + " " + company_name, "国聘网")
                    detail_link = f"https://www.iguopin.com/job/detail?id={job_id}" if job_id else ""

                    job_obj = {
                        "job_title": job_name,
                        "company": company_name,
                        "enterprise_type": enterprise_type,
                        "detail_link": detail_link,
                        "salary_range": salary_range,
                        "responsibilities": _extract_responsibilities(contents),
                        "requirements": _extract_requirements(contents, edu, exp),
                        "benefits": _generate_benefits(contents),
                        "development": _generate_development(enterprise_type, "国聘网"),
                        "work_location": work_location,
                        "source": "国聘网-城市定向（真实数据）",
                        "search_keyword": keyword,
                        "_raw_contents": contents,
                        "_raw_edu": edu,
                        "_raw_exp": exp,
                    }
                    results.append(job_obj)
                except Exception:
                    continue

            if results:
                print(f"  [国聘网-城市定向] 本次搜索返回 {len(job_list)} 条，累计 {len(results)} 条")

        except Exception as e:
            print(f"  [国聘网-城市定向] JSON解析失败: {e}")

        _random_delay()

    # 去重
    seen = set()
    deduped = []
    for r in results:
        key = (r["job_title"].lower(), r["company"].lower())
        if key not in seen:
            seen.add(key)
            deduped.append(r)

    print(f"[国聘网-城市定向] 共获取 {len(deduped)} 条「{city}」岗位")
    return deduped


# ============================================================
# 城市筛选
# ============================================================

def filter_by_city(jobs, city, min_count=5):
    """
    ★ 严格城市筛选：只保留用户常驻城市的岗位。

    规则：
    - 只返回 work_location 匹配用户城市的岗位
    - 如果匹配岗位 < min_count，仍然只返回城市匹配的（不补充全国岗位）
    - 如果城市匹配为 0，打印警告但仍返回空列表（不放宽要求）
    """
    if not city:
        return jobs

    city_jobs = [j for j in jobs if _is_city_match(j.get("work_location", ""), city)]
    other_jobs = [j for j in jobs if not _is_city_match(j.get("work_location", ""), city)]

    print(f"  严格城市筛选: 匹配「{city}」的岗位 {len(city_jobs)} 个，其他城市 {len(other_jobs)} 个（已排除）")

    if len(city_jobs) == 0:
        print(f"  ⚠️ 警告: 没有找到{city}的岗位！将不返回任何结果（不放宽城市要求）")
    elif len(city_jobs) < min_count:
        print(f"  ⚠️ {city}岗位仅 {len(city_jobs)} 个（不足{min_count}个），但仍只保留{city}岗位")

    return city_jobs


def generate_recommendations(user):
    """
    根据用户信息从多个招聘平台爬取真实在招岗位。

    参数:
        user (dict): 用户信息字典，包含:
            - id: 用户ID
            - city: 期望工作城市
            - degree: 学历（本科/硕士/博士）
            - experience: 工作经验（如"3年"）
            - field: 求职方向/关键词
            - certifications: 权威证书（可能为空）
            - email: 邮箱

    返回:
        list[dict]: 岗位推荐列表，每个字典包含:
            - job_title: 岗位名称
            - company: 公司名称
            - enterprise_type: 企业类型（国企/私企/外企）
            - match_score: 匹配度 0-100
            - detail_link: 详情链接
            - salary_range: 薪资范围
            - responsibilities: 岗位职责
            - requirements: 任职要求
            - benefits: 福利待遇
            - development: 职业发展前景
            - work_location: 工作地点
            - source: 数据来源标注
            - search_keyword: 搜索关键词
    """
    print("=" * 60)
    print(f"开始生成岗位推荐 | 用户ID: {user.get('id', 'unknown')}")
    print(f"城市: {user.get('city', '未知')} | 学历: {user.get('degree', '未知')} | "
          f"经验: {user.get('experience', '未知')} | 方向: {user.get('field', '未知')}")
    print("=" * 60)

    keyword = user.get("field") or user.get("direction") or ""
    city = user.get("city") or ""

    if not keyword:
        print("[ERROR] 缺少求职方向关键词")
        return []

    if not city:
        print("[WARNING] 未提供常驻城市，无法进行城市定向搜索")

    print(f"\n★★★ 策略: 以「{city}」常驻城市为第一优先，专业方向为第二优先 ★★★")

    all_jobs = []

    # ------------------------------------------------
    # 1. 国聘网-城市定向搜索（最高优先级，精准匹配城市）
    # ------------------------------------------------
    try:
        print(f"\n>>> [1/7] 爬取国聘网（{city}城市定向）...")
        iguopin_city_jobs = crawl_iguopin_by_city(keyword, city)
        all_jobs.extend(iguopin_city_jobs)
        print(f"国聘网-城市定向返回 {len(iguopin_city_jobs)} 条，累计 {len(all_jobs)} 条")
    except Exception as e:
        print(f"[国聘网-城市定向] 爬取失败: {e}")

    # ------------------------------------------------
    # 2. 国聘网通用搜索（补充数据）
    # ------------------------------------------------
    try:
        print("\n>>> [2/7] 爬取国聘网（通用搜索）...")
        _random_delay()
        iguopin_jobs = crawl_iguopin(keyword, city)
        all_jobs.extend(iguopin_jobs)
        print(f"国聘网通用返回 {len(iguopin_jobs)} 条，累计 {len(all_jobs)} 条")
    except Exception as e:
        print(f"[国聘网] 爬取失败: {e}")

    # ------------------------------------------------
    # 3. 央企招聘官网（中石油/南方电网/国家电网等）
    # ------------------------------------------------
    try:
        print("\n>>> [3/7] 爬取央企招聘官网（中石油/南方电网/国家电网等）...")
        _random_delay()
        soe_jobs = crawl_national_soe(keyword, city)
        all_jobs.extend(soe_jobs)
        print(f"央企官网返回 {len(soe_jobs)} 条，累计 {len(all_jobs)} 条")
    except Exception as e:
        print(f"[央企官网] 爬取失败: {e}")

    # ------------------------------------------------
    # 4. 地方国企（按用户城市匹配，如长沙银行/湖南银行）
    # ------------------------------------------------
    try:
        print(f"\n>>> [4/7] 爬取{city}地方国企招聘...")
        _random_delay()
        local_soe_jobs = crawl_local_soe(keyword, city)
        all_jobs.extend(local_soe_jobs)
        print(f"地方国企返回 {len(local_soe_jobs)} 条，累计 {len(all_jobs)} 条")
    except Exception as e:
        print(f"[地方国企] 爬取失败: {e}")

    # ------------------------------------------------
    # 5. 微信公众号招聘文章（搜狗微信搜索）
    # ------------------------------------------------
    try:
        print("\n>>> [5/7] 爬取微信公众号招聘文章（搜狗微信搜索）...")
        _random_delay()
        wechat_jobs = crawl_wechat_sogou(keyword, city)
        all_jobs.extend(wechat_jobs)
        print(f"微信公众号返回 {len(wechat_jobs)} 条，累计 {len(all_jobs)} 条")
    except Exception as e:
        print(f"[微信公众号] 爬取失败: {e}")

    # ------------------------------------------------
    # 6. BOSS直聘（Playwright渲染）
    # ------------------------------------------------
    try:
        print("\n>>> [6/7] 爬取BOSS直聘...")
        _random_delay()
        boss_jobs = crawl_boss(keyword, city)
        all_jobs.extend(boss_jobs)
        print(f"BOSS直聘返回 {len(boss_jobs)} 条，累计 {len(all_jobs)} 条")
    except Exception as e:
        print(f"[BOSS直聘] 爬取失败: {e}")

    # ------------------------------------------------
    # 7. 智联招聘（Playwright渲染）
    # ------------------------------------------------
    try:
        print("\n>>> [7/7] 爬取智联招聘...")
        _random_delay()
        zhaopin_jobs = crawl_zhaopin(keyword, city)
        all_jobs.extend(zhaopin_jobs)
        print(f"智联招聘返回 {len(zhaopin_jobs)} 条，累计 {len(all_jobs)} 条")
    except Exception as e:
        print(f"[智联招聘] 爬取失败: {e}")

    # ------------------------------------------------
    # 去重
    # ------------------------------------------------
    print(f"\n去重前: {len(all_jobs)} 条")
    seen = set()
    deduped = []
    for job in all_jobs:
        key = (job.get("job_title", "").lower(), job.get("company", "").lower())
        if key not in seen:
            seen.add(key)
            deduped.append(job)
    print(f"去重后: {len(deduped)} 条")

    # ------------------------------------------------
    # 城市筛选
    # ------------------------------------------------
    print("\n按城市筛选...")
    filtered = filter_by_city(deduped, city, min_count=5)

    # ------------------------------------------------
    # 计算匹配度
    # ------------------------------------------------
    print("\n计算匹配度...")
    for job in filtered:
        job["match_score"] = calculate_match_score(job, user)

    # 按匹配度排序
    filtered.sort(key=lambda x: x.get("match_score", 0), reverse=True)

    # ------------------------------------------------
    # 清理临时字段
    # ------------------------------------------------
    for job in filtered:
        job.pop("_raw_contents", None)
        job.pop("_raw_edu", None)
        job.pop("_raw_exp", None)
        job["search_links"] = build_search_links(
            keyword=job.get("search_keyword", ""),
            city=city,
            enterprise_type=job.get("enterprise_type", ""),
            job_title=job.get("job_title", ""),
        )

    # 按匹配度过滤：>=65% 才推荐，最多10条
    MIN_MATCH_SCORE = 65
    MAX_RESULTS = 10
    qualified = [j for j in filtered if j.get("match_score", 0) >= MIN_MATCH_SCORE]

    result = qualified[:MAX_RESULTS]

    print("\n" + "=" * 60)
    print(f"岗位推荐生成完成！共 {len(result)} 条岗位")
    print(f"来源统计:")
    source_count = {}
    for job in result:
        src = job.get("source", "未知")
        source_count[src] = source_count.get(src, 0) + 1
    for src, cnt in source_count.items():
        print(f"  {src}: {cnt} 条")
    print("=" * 60)

    return result
PUSHPLUS_TOKEN = os.environ.get("PUSHPLUS_TOKEN")
SMTP_HOST = os.environ.get("SMTP_HOST")
SMTP_PORT = int(os.environ.get("SMTP_PORT", "465"))
SMTP_USER = os.environ.get("SMTP_USER")
SMTP_PASS = os.environ.get("SMTP_PASS")

# 输出目录（GitHub Actions 中为仓库根目录）
OUTPUT_DIR = os.environ.get("GITHUB_WORKSPACE", os.path.dirname(os.path.abspath(__file__)) + "/..")
OUTPUT_DIR = os.path.abspath(OUTPUT_DIR)
REPORT_DIR = os.path.join(OUTPUT_DIR, "reports")
os.makedirs(REPORT_DIR, exist_ok=True)

print("=" * 50)
print("[环境检查]")
print(f"  SUPABASE_URL 已设置: {bool(SUPABASE_URL)}")
print(f"  SUPABASE_KEY 已设置: {bool(SUPABASE_KEY)}")
print(f"  PUSHPLUS_TOKEN 已设置: {bool(PUSHPLUS_TOKEN)}")
print(f"  SMTP 邮件推送: {'已配置' if all([SMTP_HOST, SMTP_USER, SMTP_PASS]) else '未配置'}")
print(f"  报告输出目录: {REPORT_DIR}")
print("=" * 50)

if not SUPABASE_URL:
    print("[错误] 缺少环境变量 SUPABASE_URL，请在 GitHub Secrets 中添加")
    sys.exit(1)
if not SUPABASE_KEY:
    print("[错误] 缺少环境变量 SUPABASE_KEY，请在 GitHub Secrets 中添加")
    sys.exit(1)

# Supabase REST API 请求头（使用 service_role key 绕过 RLS）
HEADERS = {
    "apikey": SUPABASE_KEY,
    "Authorization": f"Bearer {SUPABASE_KEY}",
    "Content-Type": "application/json",
    "Prefer": "return=representation",
}


# ============ 多平台真实招聘搜索链接 ============
# 国企官方招聘平台映射（根据岗位名称匹配）
ENTERPRISE_OFFICIAL_LINKS = {
    "国家电网": {"name": "国家电网招聘", "url": "https://zhaopin.sgcc.com.cn/", "icon": "⚡"},
    "南方电网": {"name": "南方电网招聘", "url": "http://www.csg.cn/", "icon": "⚡"},
    "中国移动": {"name": "中国移动招聘", "url": "https://job.10086.cn/", "icon": "📱"},
    "中国联通": {"name": "中国联通招聘", "url": "https://hr.chinaunicom.com/", "icon": "📱"},
    "中国电信": {"name": "中国电信招聘", "url": "https://www.chinatelecom.com.cn/", "icon": "📱"},
    "工商银行": {"name": "工商银行招聘", "url": "https://job.icbc.com.cn/", "icon": "🏦"},
    "建设银行": {"name": "建设银行招聘", "url": "https://job.ccb.com/", "icon": "🏦"},
    "农业银行": {"name": "农业银行招聘", "url": "https://job.abchina.com/", "icon": "🏦"},
    "中国银行": {"name": "中国银行招聘", "url": "https://campus.chinahr.com/pages/boc/", "icon": "🏦"},
    "交通银行": {"name": "交通银行招聘", "url": "https://job.bankcomm.com/", "icon": "🏦"},
    "中石油": {"name": "中石油招聘", "url": "https://zhaopin.cnpc.com.cn/", "icon": "🛢️"},
    "中石化": {"name": "中石化招聘", "url": "http://job.sinopec.com/", "icon": "🛢️"},
    "中国建筑": {"name": "中国建筑招聘", "url": "https://jobs.cscec.com/", "icon": "🏗️"},
    "中铁": {"name": "中国铁路人才招聘", "url": "https://job.crec.cn/", "icon": "🚄"},
    "中交": {"name": "中交集团招聘", "url": "https://job.ccccltd.cn/", "icon": "🚧"},
    "烟草": {"name": "中国烟草招聘", "url": "http://www.tobacco.gov.cn/", "icon": "🚬"},
    "水务": {"name": "水务集团招聘", "url": "https://www.waterchina.com/", "icon": "💧"},
    "城投": {"name": "城投集团招聘", "url": "http://www.citiccapital.com/", "icon": "🏢"},
    "地铁": {"name": "地铁集团招聘", "url": "https://www.chinametro.net/", "icon": "🚇"},
}


def get_enterprise_official_link(job_title: str) -> dict | None:
    """根据岗位名称中的国企单位名，返回对应的官方招聘平台链接"""
    for key, info in ENTERPRISE_OFFICIAL_LINKS.items():
        if key in str(job_title):
            return info
    return None


def build_search_links(keyword: str, city: str, enterprise_type: str = "", job_title: str = "") -> list[dict]:
    """为每个岗位生成多个真实招聘平台的搜索链接"""
    # 简化关键词：去掉国企前缀，只保留岗位关键词
    simple_kw = keyword
    if "-" in simple_kw:
        simple_kw = simple_kw.split("-", 1)[1].strip()
    core_kw = simple_kw
    for suffix in ["工程师", "经理", "主管", "专员", "技术员", "操作员"]:
        if core_kw.endswith(suffix) and len(core_kw) > len(suffix) + 2:
            core_kw = core_kw[:-len(suffix)]
            break
    search_kw = quote(str(core_kw or simple_kw or keyword))
    links = []
    # 综合招聘平台
    links.append({"name": "BOSS直聘", "url": f"https://www.zhipin.com/web/geek/job?query={search_kw}", "icon": "💼"})
    links.append({"name": "智联招聘", "url": f"https://sou.zhaopin.com/?kw={search_kw}", "icon": "🔍"})
    links.append({"name": "猎聘", "url": f"https://www.liepin.com/zhaopin/?key={search_kw}", "icon": "🎯"})
    # 国企官方权威平台（优先推荐）
    if enterprise_type == "国企":
        links.append({"name": "国聘网", "url": f"https://www.iguopin.com/job/list?keywords={search_kw}", "icon": "🏛️"})
        links.append({"name": "国资委央企招聘", "url": "http://www.sasac.gov.cn/n2588035/n2588105/index.html", "icon": "🇨🇳"})
        links.append({"name": "人社部就业", "url": "http://job.mohrss.gov.cn/", "icon": "📋"})
        links.append({"name": "新职业网", "url": f"https://www.ncss.cn/student/jobsearch/fairs.html?keywords={search_kw}", "icon": "🎓"})
        # 企业自有官方招聘平台
        official = get_enterprise_official_link(job_title)
        if official:
            links.append(official)
    return links

# ============ H5/docx/邮件/推送 ============
def generate_h5_report(user: dict, jobs: list[dict]) -> str:
    """生成自包含的 H5 页面，手机浏览器直接打开查看，返回文件路径"""
    user_id = user.get('id', 'unknown')
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    def esc(v):
        return html.escape(str(v)) if v else '—'

    def list_items(val):
        if isinstance(val, str):
            items = [x.strip() for x in val.split('\n') if x.strip()]
        elif isinstance(val, list):
            items = val
        else:
            items = []
        return ''.join(f'<li>{html.escape(str(x))}</li>' for x in items)

    type_colors = {
        "国企": "#dc2626",
        "私企": "#2563eb",
        "外企": "#16a34a",
    }

    cards = []
    for idx, job in enumerate(jobs, 1):
        etype = job.get('enterprise_type', '—')
        etype_color = type_colors.get(etype, "#6b7280")
        cards.append(f"""
    <div class="card">
      <div class="card-header">
        <span class="job-index">{idx}</span>
        <div class="job-title-wrap">
          <h2>{esc(job.get('job_title'))}</h2>
          <span class="tag" style="background:{etype_color}">{esc(etype)}</span>
          {f'<span class="tag source-tag">{esc(job.get("source", "AI推荐"))}</span>' if job.get('source') else ''}
        </div>
      </div>
      <div class="card-body">
        <div class="info-row">
          <div class="info-item"><span class="label">公司</span><span class="value">{esc(job.get('company'))}</span></div>
          <div class="info-item"><span class="label">薪资</span><span class="value salary">{esc(job.get('salary_range', '面议'))}</span></div>
        </div>
        <div class="info-row">
          <div class="info-item"><span class="label">匹配度</span><span class="value match">{esc(job.get('match_score', 0))}%</span></div>
          <div class="info-item"><span class="label">地点</span><span class="value">{esc(job.get('work_location'))}</span></div>
        </div>
        <div class="match-bar"><div class="match-fill" style="width:{esc(job.get('match_score', 0))}%"></div></div>
        <div class="section"><h3>📋 岗位职责</h3><ul>{list_items(job.get('responsibilities'))}</ul></div>
        <div class="section"><h3>✅ 任职要求</h3><ul>{list_items(job.get('requirements'))}</ul></div>
        <div class="section"><h3>🎁 福利待遇</h3><ul>{list_items(job.get('benefits'))}</ul></div>
        <div class="section"><h3>📈 职业发展</h3><p>{esc(job.get('development', '暂无信息'))}</p></div>
        <div class="platform-links">
          <div class="platform-title">🔗 点击查看真实在招岗位</div>
          <div class="platform-hint">💡 提示：进入招聘平台后，请手动选择城市「{esc(user.get('city', ''))}」以获取当地岗位</div>
          <div class="platform-btns">
            {''.join(f'<a class="platform-btn" href="{esc(sl["url"])}" target="_blank">{sl["icon"]} {esc(sl["name"])}</a>' for sl in job.get('search_links', []))}
          </div>
        </div>
      </div>
    </div>""")

    html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0, maximum-scale=1.0, user-scalable=no">
  <title>智能岗位推荐报告</title>
  <style>
    *{{margin:0;padding:0;box-sizing:border-box}}
    body{{font-family:-apple-system,BlinkMacSystemFont,"Segoe UI","PingFang SC","Microsoft YaHei",sans-serif;background:#f0f2f5;color:#1f2937;line-height:1.6;max-width:640px;margin:0 auto;padding-bottom:40px}}
    .header{{background:linear-gradient(135deg,#2563eb,#7c3aed);color:#fff;padding:32px 20px 24px;text-align:center;position:sticky;top:0;z-index:10;box-shadow:0 2px 12px rgba(37,99,235,.3)}}
    .header h1{{font-size:22px;margin-bottom:6px}}
    .header .date{{font-size:12px;opacity:.85}}
    .user-info{{background:#fff;margin:12px;border-radius:12px;padding:16px;box-shadow:0 1px 4px rgba(0,0,0,.06)}}
    .user-info h2{{font-size:15px;color:#2563eb;margin-bottom:12px}}
    .user-info .grid{{display:grid;grid-template-columns:1fr 1fr;gap:10px}}
    .user-info .item{{font-size:13px}}
    .user-info .item .k{{color:#6b7280;margin-right:4px}}
    .user-info .item .v{{color:#1f2937;font-weight:600}}
    .card{{background:#fff;margin:12px;border-radius:12px;overflow:hidden;box-shadow:0 1px 4px rgba(0,0,0,.06)}}
    .card-header{{display:flex;align-items:center;gap:12px;padding:16px;background:linear-gradient(90deg,#eff6ff,#f0f9ff)}}
    .job-index{{width:28px;height:28px;border-radius:50%;background:#2563eb;color:#fff;display:flex;align-items:center;justify-content:center;font-size:14px;font-weight:700;flex-shrink:0}}
    .job-title-wrap{{flex:1;min-width:0}}
    .job-title-wrap h2{{font-size:16px;color:#1f2937}}
    .tag{{display:inline-block;padding:2px 8px;border-radius:4px;font-size:11px;color:#fff;margin-top:2px}}
    .source-tag{{background:#16a34a!important;margin-left:4px}}
    .card-body{{padding:16px}}
    .info-row{{display:flex;gap:12px;margin-bottom:8px}}
    .info-item{{flex:1;display:flex;flex-direction:column}}
    .info-item .label{{font-size:11px;color:#9ca3af}}
    .info-item .value{{font-size:14px;font-weight:600;color:#1f2937}}
    .info-item .value.salary{{color:#dc2626}}
    .info-item .value.match{{color:#16a34a}}
    .match-bar{{height:6px;background:#e5e7eb;border-radius:3px;margin:8px 0 16px;overflow:hidden}}
    .match-fill{{height:100%;background:linear-gradient(90deg,#16a34a,#22c55e);border-radius:3px;transition:width .3s}}
    .section{{margin-top:14px}}
    .section h3{{font-size:13px;color:#374151;margin-bottom:6px}}
    .section ul{{padding-left:18px}}
    .section ul li{{font-size:13px;color:#4b5563;margin-bottom:4px}}
    .section p{{font-size:13px;color:#4b5563}}
    .platform-links{{margin-top:16px;border-top:1px dashed #e5e7eb;padding-top:14px}}
    .platform-title{{font-size:13px;color:#374151;margin-bottom:8px;font-weight:600}}
    .platform-hint{{font-size:11px;color:#9ca3af;margin-bottom:10px;background:#f3f4f6;padding:6px 10px;border-radius:6px}}
    .platform-btns{{display:flex;flex-wrap:wrap;gap:8px}}
    .platform-btn{{display:inline-block;padding:8px 14px;border-radius:8px;font-size:12px;font-weight:600;text-decoration:none;color:#fff;background:linear-gradient(135deg,#2563eb,#7c3aed)}}
    .platform-btn:nth-child(4n+1){{background:linear-gradient(135deg,#2563eb,#3b82f6)}}
    .platform-btn:nth-child(4n+2){{background:linear-gradient(135deg,#16a34a,#22c55e)}}
    .platform-btn:nth-child(4n+3){{background:linear-gradient(135deg,#ea580c,#f97316)}}
    .platform-btn:nth-child(4n){{background:linear-gradient(135deg,#dc2626,#ef4444)}}
    .footer{{text-align:center;padding:20px;font-size:11px;color:#9ca3af}}
  </style>
</head>
<body>
  <div class="header">
    <h1>📋 智能岗位推荐报告</h1>
    <div class="date">{datetime.now().strftime("%Y年%m月%d日 %H:%M")} 生成</div>
  </div>
  <div class="user-info">
    <h2>👤 求职者信息</h2>
    <div class="grid">
      <div class="item"><span class="k">城市</span><span class="v">{esc(user.get('city'))}</span></div>
      <div class="item"><span class="k">学历</span><span class="v">{esc(user.get('degree'))}</span></div>
      <div class="item"><span class="k">工作经验</span><span class="v">{esc(user.get('experience'))}</span></div>
      <div class="item"><span class="k">求职方向</span><span class="v">{esc(user.get('field'))}</span></div>
      <div class="item"><span class="k">权威证书</span><span class="v">{esc(user.get('certifications') or '无')}</span></div>
    </div>
  </div>
  {''.join(cards)}
  <div class="footer">本报告由智能岗位推荐系统自动生成 | Powered by 智谱 GLM-4-Flash</div>
</body>
</html>"""

    filename = f"job_h5_user{user_id}_{timestamp}.html"
    filepath = os.path.join(REPORT_DIR, filename)
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(html_content)
    print(f"[H5] 已生成: {filepath}")
    return filepath


# ============ 生成 docx 岗位推荐报告 ============
def generate_docx_report(user: dict, jobs: list[dict]) -> str:
    """生成 docx 格式的岗位推荐报告，返回文件路径"""
    doc = Document()

    # ---- 设置默认字体 ----
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # ---- 标题 ----
    title = doc.add_heading('智能岗位推荐报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)

    # 日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(
        f'生成日期：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}'
    )
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

    doc.add_paragraph()  # 空行

    # ---- 求职者信息 ----
    doc.add_heading('一、求职者信息', level=1)
    info_table = doc.add_table(rows=3, cols=4, style='Light Shading Accent 1')
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ['城市', user.get('city', '—'), '学历', user.get('degree', '—')],
        ['工作经验', user.get('experience', '—'), '求职方向', user.get('field', '—')],
        ['权威证书', user.get('certifications') or '无', '邮箱', user.get('email') or '未提供'],
    ]
    for row_idx, row_data in enumerate(info_data):
        for col_idx, val in enumerate(row_data):
            cell = info_table.cell(row_idx, col_idx)
            cell.text = val
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if col_idx % 2 == 0:  # 标签列加粗
                        run.bold = True
                        run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)

    doc.add_paragraph()  # 空行

    # ---- 推荐岗位详情 ----
    doc.add_heading('二、推荐岗位详情', level=1)

    type_colors = {
        "国企": RGBColor(0xdc, 0x26, 0x26),
        "私企": RGBColor(0x25, 0x63, 0xeb),
        "外企": RGBColor(0x16, 0xa3, 0x4a),
    }

    for idx, job in enumerate(jobs, 1):
        # 岗位标题
        heading = doc.add_heading(f'{idx}. {job["job_title"]}', level=2)

        # 基本信息（表格）
        basic_table = doc.add_table(rows=2, cols=4, style='Light List Accent 1')
        basic_data = [
            ['公司名称', job.get('company', '—'), '企业类型', job.get('enterprise_type', '—')],
            ['匹配度', f'{job.get("match_score", 0)}%', '薪资范围', job.get('salary_range', '面议')],
        ]
        for row_idx, row_data in enumerate(basic_data):
            for col_idx, val in enumerate(row_data):
                cell = basic_table.cell(row_idx, col_idx)
                cell.text = val
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
                        if col_idx % 2 == 0:
                            run.bold = True

        # 工作地点
        if job.get('work_location'):
            loc_para = doc.add_paragraph()
            loc_label = loc_para.add_run('工作地点：')
            loc_label.bold = True
            loc_label.font.size = Pt(10)
            loc_val = loc_para.add_run(job['work_location'])
            loc_val.font.size = Pt(10)

        # 岗位职责
        doc.add_heading('岗位职责', level=3)
        responsibilities = job.get('responsibilities', '')
        if isinstance(responsibilities, str):
            resp_list = [r.strip() for r in responsibilities.split('\n') if r.strip()]
        else:
            resp_list = responsibilities
        for item in resp_list:
            p = doc.add_paragraph(item, style='List Bullet')
            for run in p.runs:
                run.font.size = Pt(10)

        # 任职要求
        doc.add_heading('任职要求', level=3)
        requirements = job.get('requirements', '')
        if isinstance(requirements, str):
            req_list = [r.strip() for r in requirements.split('\n') if r.strip()]
        else:
            req_list = requirements
        for item in req_list:
            p = doc.add_paragraph(item, style='List Bullet')
            for run in p.runs:
                run.font.size = Pt(10)

        # 福利待遇
        doc.add_heading('福利待遇', level=3)
        benefits = job.get('benefits', '')
        if isinstance(benefits, str):
            ben_list = [r.strip() for r in benefits.split('\n') if r.strip()]
        else:
            ben_list = benefits
        for item in ben_list:
            p = doc.add_paragraph(item, style='List Bullet')
            for run in p.runs:
                run.font.size = Pt(10)

        # 职业发展前景
        doc.add_heading('职业发展前景', level=3)
        dev_para = doc.add_paragraph(job.get('development', '暂无信息'))
        for run in dev_para.runs:
            run.font.size = Pt(10)

        # 各平台搜索链接
        doc.add_heading('查看真实在招岗位', level=3)
        for sl in job.get('search_links', []):
            p = doc.add_paragraph()
            run = p.add_run(f"{sl['icon']} {sl['name']}：{sl['url']}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)

        # 分隔线（除最后一个外）
        if idx < len(jobs):
            sep = doc.add_paragraph()
            sep_run = sep.add_run('—' * 40)
            sep_run.font.color.rgb = RGBColor(0xd1, 0xd5, 0xdb)
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- 页脚 ----
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('本报告由智能岗位推荐系统自动生成 | Powered by 智谱 GLM-4-Flash')
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)

    # ---- 保存文件 ----
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    user_id = user.get('id', 'unknown')
    filename = f"job_report_user{user_id}_{timestamp}.docx"
    filepath = os.path.join(REPORT_DIR, filename)
    doc.save(filepath)
    print(f"[报告] 已生成: {filepath}")
    return filepath


# ============ 上传报告到 Supabase Storage ============
STORAGE_BUCKET = "reports"  # 需要在 Supabase 创建名为 reports 的公开 bucket


def upload_to_storage(filepath: str, content_type: str = "application/vnd.openxmlformats-officedocument.wordprocessingml.document") -> str | None:
    """上传文件到 Supabase Storage，返回公开下载链接；失败则返回 None"""
    filename = os.path.basename(filepath)
    try:
        with open(filepath, "rb") as f:
            file_data = f.read()
        upload_headers = {
            "apikey": SUPABASE_KEY,
            "Authorization": f"Bearer {SUPABASE_KEY}",
            "Content-Type": content_type,
        }
        resp = requests.post(
            f"{SUPABASE_URL}/storage/v1/object/{STORAGE_BUCKET}/{filename}",
            headers=upload_headers,
            data=file_data,
            timeout=60,
        )
        if resp.status_code in (200, 201):
            public_url = f"{SUPABASE_URL}/storage/v1/object/public/{STORAGE_BUCKET}/{filename}"
            print(f"[Storage] 上传成功，公开链接: {public_url}")
            return public_url
        else:
            print(f"[Storage] 上传失败 HTTP {resp.status_code}: {resp.text}")
            return None
    except Exception as e:
        print(f"[Storage] 上传异常: {e}")
        return None


# ============ 邮件推送 ============
def send_email(user: dict, h5_url: str | None, jobs: list[dict]):
    """通过 SMTP 发送邮件，包含 H5 链接和岗位摘要"""
    user_email = user.get('email')
    if not user_email:
        print("[邮件] 用户未提供邮箱，跳过")
        return
    if not all([SMTP_HOST, SMTP_USER, SMTP_PASS]):
        print("[邮件] SMTP 未配置，跳过邮件推送")
        return

    city = user.get('city', '—')
    field = user.get('field', '求职方向')
    subject = f"【岗位推荐】{field} - {city}（{datetime.now().strftime('%m月%d日 %H:%M')}）"

    # 岗位摘要
    job_lines = []
    for i, j in enumerate(jobs, 1):
        job_lines.append(
            f"  {i}. {j.get('job_title', '—')}"
            f" | 匹配度 {j.get('match_score', 0)}%"
            f" | 薪资 {j.get('salary_range', '面议')}"
            f" | {j.get('enterprise_type', '—')}"
        )
    jobs_summary = "\n".join(job_lines)

    body = f"""您好！

您的智能岗位推荐报告已生成。

══════════════════════════════
求职者：{city} | {field}
══════════════════════════════

本次推荐 {len(jobs)} 个岗位：
{jobs_summary}

──────────────────────────────
📱 点击查看完整岗位推荐详情（H5页面）：
{h5_url or '（链接生成失败，请稍后在GitHub Actions中查看）'}
──────────────────────────────

H5页面中包含：
  - 每个岗位的详细职责、任职要求、福利待遇
  - 各官方招聘平台的真实在招岗位搜索链接
    （国聘网/国资委/人社部/新职业网/企业官网）

祝您求职顺利！

---
智能岗位推荐系统
生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}
Powered by 智谱 GLM-4-Flash
"""

    msg = MIMEText(body, "plain", "utf-8")
    msg["Subject"] = subject
    msg["From"] = SMTP_USER
    msg["To"] = user_email

    try:
        with smtplib.SMTP_SSL(SMTP_HOST, SMTP_PORT, timeout=30) as server:
            server.login(SMTP_USER, SMTP_PASS)
            server.sendmail(SMTP_USER, [user_email], msg.as_string())
        print(f"[邮件] 已发送至 {user_email}")
    except Exception as e:
        print(f"[邮件] 发送失败: {e}")


# ============ PushPlus 微信推送 ============
def push_wechat(user: dict, jobs: list[dict], download_url: str | None = None):
    if not PUSHPLUS_TOKEN:
        return
    title = f"岗位推荐：{user.get('field', '求职方向')}"
    # 使用 Markdown 模板，内容包含完整岗位详情 + 下载链接
    md_lines = [
        f"## 📋 岗位推荐报告",
        f"",
        f"**求职者信息**",
        f"- 城市：{user.get('city', '—')}",
        f"- 学历：{user.get('degree', '—')}",
        f"- 工作经验：{user.get('experience', '—')}",
        f"- 求职方向：{user.get('field', '—')}",
        f"",
        f"---",
        f"",
    ]
    for i, j in enumerate(jobs, 1):
        md_lines.append(f"### {i}. {j['job_title']}")
        md_lines.append(f"")
        md_lines.append(f"| 项目 | 详情 |")
        md_lines.append(f"|------|------|")
        md_lines.append(f"| 公司 | {j.get('company', '—')}（{j.get('enterprise_type', '—')}）|")
        md_lines.append(f"| 匹配度 | **{j.get('match_score', 0)}%** |")
        md_lines.append(f"| 薪资 | {j.get('salary_range', '面议')} |")
        if j.get('work_location'):
            md_lines.append(f"| 地点 | {j['work_location']} |")
        md_lines.append(f"| 来源 | {j.get('source', 'AI推荐方向')} |")
        md_lines.append(f"")
        # 各平台搜索链接
        if j.get('search_links'):
            md_lines.append(f"**🔗 查看真实在招岗位：**")
            for sl in j['search_links']:
                md_lines.append(f"- [{sl['icon']} {sl['name']}]({sl['url']})")
            md_lines.append(f"")
        md_lines.append(f"---")
        md_lines.append(f"")
    # 查看链接（H5 页面，手机直接打开）
    if download_url:
        md_lines.append(f"📱 **[点击查看完整岗位推荐详情]({download_url})**")
        md_lines.append(f"")
        md_lines.append(f"> 手机浏览器直接打开，查看每个岗位的职责、要求、福利等完整信息")
    else:
        md_lines.append(f"> 完整岗位推荐报告已生成，请前往 GitHub Actions 下载")
    content = "\n".join(md_lines)
    try:
        resp = requests.post(
            "http://www.pushplus.plus/send",
            json={
                "token": PUSHPLUS_TOKEN,
                "title": title,
                "content": content,
                "template": "markdown",
            },
            timeout=15,
        )
        if resp.status_code == 200:
            print(f"[PushPlus] 推送成功")
        else:
            print(f"[PushPlus] 推送返回 HTTP {resp.status_code}: {resp.text}")
    except Exception as e:
        print(f"[PushPlus] 推送失败：{e}")


# ============ 主流程 ============
def _is_garbled(text):
    """检测文本是否乱码（含不可打印字符或过量问号）"""
    if not text:
        return True
    # 全是问号或替换字符
    if text.strip().replace("?", "").replace("�", "").replace(" ", "").replace("\ufffd", "").strip() == "":
        return True
    # 问号占比超过30%
    q_count = text.count("?") + text.count("�") + text.count("\ufffd")
    if q_count > len(text) * 0.3:
        return True
    return False

def main():
    MAX_USERS = 12  # 每批最多处理12个用户（约60分钟）
    try:
        users = fetch_users()
        print(f"[主流程] 共读取到 {len(users)} 个用户")

        if len(users) == 0:
            print("[主流程] 没有用户数据，任务结束")
            return

        # 过滤乱码用户
        valid_users = []
        for u in users:
            field = str(u.get("field", ""))
            city = str(u.get("city", ""))
            if _is_garbled(field) or _is_garbled(city):
                print(f"[主流程] ⚠️ 跳过乱码用户 {u.get('id')}: field='{field}' city='{city}'")
                continue
            valid_users.append(u)
        
        print(f"[主流程] 有效用户 {len(valid_users)} 个，跳过 {len(users) - len(valid_users)} 个乱码")
        
        # 限制每批数量
        batch = valid_users[:MAX_USERS]
        if len(batch) < len(valid_users):
            print(f"[主流程] 限制本批处理 {MAX_USERS} 个用户，剩余 {len(valid_users) - MAX_USERS} 个留待下次")

        success_count = 0
        fail_count = 0
        generated_reports = []

        for u in batch:
            try:
                jobs = generate_recommendations(u)
                save_recommendations(u["id"], jobs)
                # 先生成 H5 页面并上传，获取可直接跳转的链接
                h5_path = generate_h5_report(u, jobs)
                h5_url = upload_to_storage(h5_path, content_type="text/html; charset=utf-8")
                # 将每个岗位的详情链接指向 H5 页面
                if h5_url:
                    for j in jobs:
                        j["detail_link"] = h5_url
                # 生成 docx 报告
                report_path = generate_docx_report(u, jobs)
                generated_reports.append(report_path)
                push_wechat(u, jobs, h5_url)
                send_email(u, h5_url, jobs)
                print(f"[主流程] 用户 {u['id']} 推荐了 {len(jobs)} 个岗位，H5+docx+邮件已推送")
                success_count += 1
            except Exception as e:
                print(f"[主流程] 用户 {u.get('id')} 处理失败: {e}")
                traceback.print_exc()
                fail_count += 1

        print(f"[主流程] 完成！成功 {success_count} 个，失败 {fail_count} 个")
        print(f"[报告] 共生成 {len(generated_reports)} 份 docx 报告:")
        for r in generated_reports:
            print(f"  - {r}")

    except Exception as e:
        print(f"[主流程] 致命错误: {e}")
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
